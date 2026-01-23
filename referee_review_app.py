#!/usr/bin/env python3
"""
Rugby Referee Review System
Professional review application based on CRRDF framework

Copyright © 2025 Andrew Clarkson
All Rights Reserved

This application implements the Community Rugby Referee Development Framework (CRRDF).

Licensed under MIT License

Version: 2.1.0-alpha2
"""

__version__ = "2.1.0-alpha2"
__author__ = "Andrew Clarkson"
__copyright__ = "Copyright © 2025 Andrew Clarkson"
__license__ = "MIT"

# Application Constants
CRRDF_VERSION = "February 2025"

# Common Rugby Game Grades
GAME_GRADES = [
    "Division 1",
    "Southern Premier",
    "Central Premier",
    "Women's Division 1",
    "Premier 1st XV",
    "Division 2",
    "Division 3",
    "U20 Colts",
    "U21 Colts Division 1",
    "U21 Colts Division 2",
    "Women's Division 2",
    "High School",
    "Other"
]

import tkinter as tk
from tkinter import messagebox, filedialog
import sys
import os
try:
    import ttkbootstrap as ttk
    from ttkbootstrap.constants import *
    THEME_AVAILABLE = True
except ImportError:
    from tkinter import ttk
    THEME_AVAILABLE = False
    print("Warning: ttkbootstrap not available, using standard theme")

import json
from datetime import datetime
from pathlib import Path
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from typing import Dict, List, Any
import matplotlib
matplotlib.use('TkAgg')  # Use Tk backend
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure

# CRRDF Framework Questions Database
CRRDF_QUESTIONS = {
    "technical": {
        "title": "Technical Pillar",
        "questions": [
            {
                "q": "How well did you apply the laws to keep the game safe?",
                "prompts": ["Think about scrum safety", "Tackle height compliance", "Dangerous play management"]
            },
            {
                "q": "Were your decisions based on actual player actions rather than assumptions?",
                "prompts": ["Pre-contest positioning", "Contest observation", "Post-contest management"]
            },
            {
                "q": "How effectively did you apply game innovations (e.g., tackle height, scrum offside)?",
                "prompts": ["Clarity on new laws", "Consistency of application", "Communication to players"]
            }
        ]
    },
    "tactical": {
        "title": "Tactical Pillar",
        "questions": [
            {
                "q": "Did you use tactical awareness in your decision-making?",
                "prompts": ["Impact vs. relevance of infringements", "Pressure and dominance assessment", "Intentional vs. accidental actions"]
            },
            {
                "q": "How appropriate was your advantage application?",
                "prompts": ["Team skill levels considered", "Game flow maintained", "Tactical vs. territorial advantage"]
            },
            {
                "q": "Were your running lines and positioning optimal for decision-making?",
                "prompts": ["Exit points from phases", "Contestable moments coverage", "Adaptation to game shape"]
            },
            {
                "q": "Did you show appropriate empathy based on game context?",
                "prompts": ["Skill level consideration", "Team tactics awareness", "Weather/conditions impact"]
            }
        ]
    },
    "management": {
        "title": "Management Pillar",
        "questions": [
            {
                "q": "Did you recognize and apply the right level of management for this game?",
                "prompts": ["Player temperament assessment", "Law knowledge and compliance levels", "Connection with captains"]
            },
            {
                "q": "How effectively did you manage player safety?",
                "prompts": ["Recognition of dangerous play", "Proactive safety management", "Work with key players"]
            },
            {
                "q": "What range of management strategies did you use?",
                "prompts": ["Captain engagement", "Verbal management balance", "Escalation when needed"]
            },
            {
                "q": "Did you identify and address game trends?",
                "prompts": ["Team tactics recognition", "Appropriate escalation", "Individual player management"]
            }
        ]
    },
    "mental": {
        "title": "Mental Pillar",
        "questions": [
            {
                "q": "How effective was your pre-game preparation?",
                "prompts": ["Mental preparation", "Game plan development", "Kit and logistics organization"]
            },
            {
                "q": "Did you identify and mitigate any 'clutter' affecting performance?",
                "prompts": ["External pressures recognized", "Strategies used to minimize impact", "Focus maintenance"]
            },
            {
                "q": "What key soft skills did you demonstrate?",
                "prompts": ["Fortitude in big moments", "Adaptability to game changes", "Resilience after challenges", "Presence and character"]
            }
        ]
    },
    "physical": {
        "title": "Physical Pillar",
        "questions": [
            {
                "q": "Did you keep up with play throughout the game?",
                "prompts": ["Fitness sufficiency", "Positioning for key moments", "Energy management"]
            }
        ]
    }
}

# GFA Categories with detailed descriptions
GFA_CATEGORIES = {
    "Tackle/Ruck": [
        ("Tackler", "Tackler obligations: release of ball carrier, clear release before playing ball, not sealing off"),
        ("Jackler", "Jackler at breakdown: on feet, clear release, legal entry gate, supporting own weight")
    ],
    "Scrum": [
        ("Stability", "Scrum stability & safety: early engagement, binding, collapse management, player welfare"),
        ("Push Straight", "Scrum law application: 1.5m push, wheel management, offside (feeding halfback)")
    ],
    "LO/Maul": [
        ("Fair Contest in the air", "Lineout contest: straight throw, fair jump, no early movement, safety in the air"),
        ("Legal maul set up", "Maul: legal formation, binding, use it call, obstruction, collapsing")
    ],
    "Space": [
        ("Kicks", "Kicks in general play: 10m offside, 50:22, contestable kicks, chasers, protection of kicker"),
        ("Breakdown/Maul", "Space at breakdown/maul: entry gate, offside lines, clear ruck/maul, quick ball vs contest")
    ],
    "Management": [
        ("Advantage", "Advantage application: territorial vs tactical, team skill consideration, materiality, clear calls"),
        ("Foul Play/Big Moments", "Foul play & pressure moments: cards, penalty tries, TMO, managing tensions, big calls")
    ]
}

RATING_SCALE = {
    1: "Unacceptable - Fails to demonstrate required awareness. Errors significantly impacted the game.",
    2: "Below Standard - Limited awareness. Errors had negative influence on the game.",
    3: "Satisfactory - Adequate level. Minor errors didn't materially affect the game. Meets minimum expectations.",
    4: "Sound - Good awareness and execution. Generally accurate and timely. Contributed positively.",
    5: "Excellent - Consistently high level. Accurate, proactive, adds clear value. Sets best practice example."
}


class PlaceholderEntry(tk.Entry):
    """Entry widget with placeholder text"""
    
    def __init__(self, master=None, placeholder="", **kwargs):
        super().__init__(master, **kwargs)
        
        self.placeholder = placeholder
        self.placeholder_color = '#999999'
        self.default_fg_color = kwargs.get('fg', 'black')
        
        self.bind("<FocusIn>", self.on_focus_in)
        self.bind("<FocusOut>", self.on_focus_out)
        
        self.put_placeholder()
    
    def put_placeholder(self):
        if not self.get():
            self.insert(0, self.placeholder)
            self.config(fg=self.placeholder_color)
    
    def on_focus_in(self, event):
        if self.get() == self.placeholder:
            self.delete(0, tk.END)
            self.config(fg=self.default_fg_color)
    
    def on_focus_out(self, event):
        if not self.get():
            self.put_placeholder()
    
    def get_value(self):
        """Get the actual value (empty string if only placeholder)"""
        current = self.get()
        return "" if current == self.placeholder else current


class PlaceholderText(tk.Text):
    """Text widget with placeholder text"""
    
    def __init__(self, master=None, placeholder="", **kwargs):
        super().__init__(master, **kwargs)
        
        self.placeholder = placeholder
        self.placeholder_color = '#999999'
        self.default_fg_color = kwargs.get('fg', 'black')
        
        self.bind("<FocusIn>", self.on_focus_in)
        self.bind("<FocusOut>", self.on_focus_out)
        
        # Add tab navigation - move to next widget instead of inserting tab
        def focus_next(event):
            event.widget.tk_focusNext().focus()
            return "break"  # Prevent default tab behavior
        self.bind("<Tab>", focus_next)
        
        self.put_placeholder()
    
    def put_placeholder(self):
        if not self.get("1.0", "end-1c"):
            self.insert("1.0", self.placeholder)
            self.config(fg=self.placeholder_color)
    
    def on_focus_in(self, event):
        if self.get("1.0", "end-1c") == self.placeholder:
            self.delete("1.0", tk.END)
            self.config(fg=self.default_fg_color)
    
    def on_focus_out(self, event):
        if not self.get("1.0", "end-1c"):
            self.put_placeholder()
    
    def get_value(self):
        """Get the actual value (empty string if only placeholder)"""
        current = self.get("1.0", "end-1c")
        return "" if current == self.placeholder else current


class ReviewSession:
    """Stores all review data for a game"""
    
    def __init__(self):
        self.metadata = {
            "game_grade": "",
            "date": datetime.now().strftime("%Y-%m-%d"),
            "result": "",
            "referee": "",
            "coach": "",
            "date_completed": datetime.now().strftime("%Y-%m-%d")
        }
        self.goals = {
            "primary": "",
            "secondary": ""
        }
        self.difficulty = 5  # 1-10 scale
        self.reflections = {
            "goals_met": "",
            "what_went_well": "",
            "biggest_challenge": "",
            "taking_forward": ""
        }
        self.crrdf_reflections = {}
        self.gfa_scores = {}
        self.coach_feedback = ""


class IDPData:
    """Individual Development Plan data structure"""
    def __init__(self):
        # Meta
        self.referee_name = ""
        self.date = datetime.now().strftime("%Y-%m-%d")
        self.last_updated = datetime.now().isoformat()
        
        # Current Reality - Club
        self.club_level = ""
        self.club_description = ""
        self.club_challenges = ""
        self.club_goal = ""
        
        # Current Reality - Rep
        self.rep_involvement = ""
        self.rep_details = ""
        self.rep_goals = ""
        
        # Current Reality - Aspirations
        self.ultimate_goal = ""
        self.season_goal = ""
        self.success_criteria = ""
        
        # The How - Fitness
        self.fitness_level = ""
        self.bronco_done = ""
        self.bronco_time = ""
        self.bronco_target = ""
        self.training_frequency = ""
        self.training_types = ""
        self.fitness_target = ""
        self.fitness_obstacles = ""
        
        # The How - Law
        self.law_confidence = ""
        self.law_weak_areas = ""
        self.law_consistency = ""
        self.law_plan = ""
        
        # The How - Mental
        self.mental_strengths = ""
        self.mental_challenges = ""
        self.mental_plan = ""
        
        # Focus Areas
        self.focus1_category = ""
        self.focus1_area = ""
        self.focus1_why = ""
        self.focus1_goal = ""
        self.focus1_how = ""
        self.focus1_track = ""
        self.focus1_obstacles = ""
        
        self.focus2_category = ""
        self.focus2_area = ""
        self.focus2_why = ""
        self.focus2_goal = ""
        self.focus2_how = ""
        self.focus2_track = ""
        self.focus2_obstacles = ""
        
        self.focus3_category = ""
        self.focus3_area = ""
        self.focus3_why = ""
        self.focus3_goal = ""
        self.focus3_how = ""
        self.focus3_track = ""
        self.focus3_obstacles = ""


class CRRDFReviewApp:
    """Main application window"""
    
    def __init__(self, root):
        self.root = root
        self.root.title("Rugby Referee Review System")
        
        # Set window icon
        try:
            # Get the correct path for icon whether running as script or bundled
            if getattr(sys, 'frozen', False):
                # Running as compiled executable
                application_path = sys._MEIPASS
            else:
                # Running as script
                application_path = os.path.dirname(os.path.abspath(__file__))
            
            icon_path = os.path.join(application_path, 'icon.ico')
            if os.path.exists(icon_path):
                self.root.iconbitmap(icon_path)
        except Exception as e:
            # If icon file not found, continue without it
            pass
        
        # Maximize window on startup
        self.root.state('zoomed')  # Windows/Linux
        # For macOS, use: self.root.attributes('-zoomed', True)
        
        # Set minimum window size
        self.root.minsize(1000, 800)
        
        # Start with a good default size for home screen
        self.root.geometry("1200x850")
        
        # Center window on screen
        self.center_window()
        
        # Allow window to be resizable
        self.root.resizable(True, True)
        
        # Bind keyboard shortcuts
        self.setup_keyboard_shortcuts()
        
        self.session = ReviewSession()
        self.current_pillar = None
        self.current_question = 0
        
        # Set up reviews directory FIRST (needed by IDP)
        self.reviews_dir = Path.home() / "Documents" / "RugbyRefereeReviews"
        self.reviews_dir.mkdir(parents=True, exist_ok=True)
        
        # Set up config file
        self.config_file = Path.home() / ".rugby_referee_review_config.json"
        self.config = self.load_config()
        
        # IDP data (after reviews_dir is set)
        self.idp_data = None
        self.idp_current_section = 0
        self.idp_sections = [
            "club", "rep", "aspirations",
            "fitness", "law", "mental",
            "focus1", "focus2", "focus3"
        ]
        self.load_idp()
        
        # Check if first run
        if self.config.get('first_run', True):
            self.show_welcome_screen()
            return  # Don't continue with normal UI until welcome is done
        
        # Apply modern theme if available
        if THEME_AVAILABLE:
            style = ttk.Style("cosmo")  # Modern, professional theme
        else:
            style = ttk.Style()
            style.theme_use('clam')
        
        # Create menu bar
        self.create_menu_bar()
        
        # Create main container and show home screen
        self.create_widgets()
        self.show_home_screen()
    
    def get_review_filename(self, referee_name, date):
        """Generate standardized filename for review"""
        # Sanitize filename
        safe_name = "".join(c for c in referee_name if c.isalnum() or c in (' ', '-', '_')).strip()
        if not safe_name:
            safe_name = "Unknown"
        return self.reviews_dir / f"Review_{safe_name}_{date}.json"
    
    def load_config(self):
        """Load configuration from file"""
        try:
            if self.config_file.exists():
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except:
            pass
        return {"first_run": True}
    
    def save_config(self):
        """Save configuration to file"""
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"Failed to save config: {e}")
    
    def load_idp(self):
        """Load IDP from file"""
        idp_file = self.reviews_dir / "idp.json"
        try:
            if idp_file.exists():
                with open(idp_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    idp = IDPData()
                    # Load all fields
                    for key, value in data.items():
                        if hasattr(idp, key):
                            setattr(idp, key, value)
                    self.idp_data = idp
                    return True
        except Exception as e:
            print(f"Failed to load IDP: {e}")
        return False
    
    def save_idp(self):
        """Save IDP to file"""
        if not self.idp_data:
            return False
        
        idp_file = self.reviews_dir / "idp.json"
        try:
            # Update last modified
            self.idp_data.last_updated = datetime.now().isoformat()
            
            # Convert to dict
            data = {}
            for key in dir(self.idp_data):
                if not key.startswith('_'):
                    value = getattr(self.idp_data, key)
                    if not callable(value):
                        data[key] = value
            
            with open(idp_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            return True
        except Exception as e:
            messagebox.showerror("Save Error", f"Failed to save IDP:\n{str(e)}")
            return False
    
    def show_welcome_screen(self):
        """Show first-run welcome and setup screen"""
        # Clear any existing content
        for widget in self.root.winfo_children():
            widget.destroy()
        
        # Create welcome screen
        welcome = tk.Frame(self.root, bg="white")
        welcome.pack(fill=tk.BOTH, expand=True)
        
        # Header
        header = tk.Frame(welcome, bg="#1976D2", height=100)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        tk.Label(header, text="Welcome to", 
                font=("Segoe UI", 14), bg="#1976D2", fg="#E3F2FD").pack(pady=(20, 0))
        tk.Label(header, text="Rugby Referee Review System", 
                font=("Segoe UI", 24, "bold"), bg="#1976D2", fg="white").pack(pady=(5, 20))
        
        # Content
        content = tk.Frame(welcome, bg="white")
        content.pack(fill=tk.BOTH, expand=True, padx=50, pady=40)
        
        tk.Label(content, text="Let's personalize your experience", 
                font=("Segoe UI", 16, "bold"), bg="white", fg="#212121").pack(pady=(0, 10))
        
        tk.Label(content, text="This information will be used to auto-fill your review forms.", 
                font=("Segoe UI", 10), bg="white", fg="#757575").pack(pady=(0, 30))
        
        # Name field
        name_frame = tk.Frame(content, bg="white")
        name_frame.pack(fill=tk.X, pady=10)
        
        tk.Label(name_frame, text="Your Full Name:", 
                font=("Segoe UI", 11, "bold"), bg="white", fg="#212121", anchor="w").pack(fill=tk.X, pady=(0, 5))
        tk.Label(name_frame, text="This will auto-fill the 'Referee' field in your reviews", 
                font=("Segoe UI", 9), bg="white", fg="#757575", anchor="w").pack(fill=tk.X, pady=(0, 5))
        
        name_entry = tk.Entry(name_frame, font=("Segoe UI", 12), width=40)
        name_entry.pack(fill=tk.X, pady=5)
        name_entry.focus()
        
        # Coach field
        coach_frame = tk.Frame(content, bg="white")
        coach_frame.pack(fill=tk.X, pady=20)
        
        tk.Label(coach_frame, text="Preferred Coach Name (Optional):", 
                font=("Segoe UI", 11, "bold"), bg="white", fg="#212121", anchor="w").pack(fill=tk.X, pady=(0, 5))
        tk.Label(coach_frame, text="This will auto-fill the 'Coach' field in your reviews", 
                font=("Segoe UI", 9), bg="white", fg="#757575", anchor="w").pack(fill=tk.X, pady=(0, 5))
        
        coach_entry = tk.Entry(coach_frame, font=("Segoe UI", 12), width=40)
        coach_entry.pack(fill=tk.X, pady=5)
        
        # Error label
        error_label = tk.Label(content, text="", font=("Segoe UI", 10), 
                              bg="white", fg="#F44336")
        error_label.pack(pady=10)
        
        def complete_setup():
            """Save settings and continue"""
            name = name_entry.get().strip()
            
            if not name:
                error_label.config(text="Please enter your name")
                return
            
            coach = coach_entry.get().strip()
            
            # Save config
            self.config = {
                "first_run": False,
                "user_name": name,
                "coach_name": coach
            }
            self.save_config()
            
            # Clear welcome screen
            welcome.destroy()
            
            # Initialize normal UI
            if THEME_AVAILABLE:
                style = ttk.Style("cosmo")
            else:
                style = ttk.Style()
                style.theme_use('clam')
            
            self.create_menu_bar()
            self.create_widgets()
            self.show_home_screen()
        
        # Buttons
        button_frame = tk.Frame(content, bg="white")
        button_frame.pack(pady=30)
        
        tk.Button(button_frame, text="Get Started", command=complete_setup,
                 font=("Segoe UI", 12, "bold"), bg="#4CAF50", fg="white",
                 padx=40, pady=12, relief=tk.FLAT, cursor="hand2").pack()
        
        # Bind Enter key
        name_entry.bind('<Return>', lambda e: complete_setup())
        coach_entry.bind('<Return>', lambda e: complete_setup())
    
    def save_review_json(self):
        """Save current review to JSON file"""
        self.status_saving()
        try:
            # Compile all review data
            review_data = {
                "version": __version__,
                "saved_at": datetime.now().isoformat(),
                "metadata": self.session.metadata,
                "goals": self.session.goals,
                "difficulty": self.session.difficulty,
                "reflections": self.session.reflections,
                "crrdf_reflections": self.session.crrdf_reflections if hasattr(self.session, 'crrdf_reflections') else {},
                "gfa_scores": self.session.gfa_scores,
                "coach_feedback": self.session.coach_feedback
            }
            
            # Generate filename
            filename = self.get_review_filename(
                self.session.metadata.get('referee', 'Unknown'),
                self.session.metadata.get('date', datetime.now().strftime("%Y-%m-%d"))
            )
            
            # Save to JSON
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(review_data, f, indent=2, ensure_ascii=False)
            
            return filename
        except Exception as e:
            print(f"Failed to save JSON: {e}")
            return None
    
    def load_review_json(self, filename):
        """Load review from JSON file"""
        try:
            with open(filename, 'r', encoding='utf-8') as f:
                review_data = json.load(f)
            
            # Load into session
            self.session = ReviewSession()
            self.session.metadata = review_data.get('metadata', {})
            self.session.goals = review_data.get('goals', {})
            self.session.difficulty = review_data.get('difficulty', 5)
            self.session.reflections = review_data.get('reflections', {})
            self.session.crrdf_reflections = review_data.get('crrdf_reflections', {})
            self.session.gfa_scores = review_data.get('gfa_scores', {})
            self.session.coach_feedback = review_data.get('coach_feedback', "")
            
            # Also populate pillar_answers for display
            self.pillar_answers = review_data.get('crrdf_reflections', {}).copy()
            
            return True
        except Exception as e:
            messagebox.showerror("Load Error", f"Failed to load review:\n{str(e)}")
            return False
    
    def center_window(self):
        """Center the window on screen"""
        self.root.update_idletasks()
        width = 1100
        height = 750
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')
    
    def setup_keyboard_shortcuts(self):
        """Setup keyboard shortcuts for common actions"""
        # Ctrl+H - Home
        self.root.bind('<Control-h>', lambda e: self.show_home_screen())
        self.root.bind('<Control-H>', lambda e: self.show_home_screen())
        
        # Ctrl+N - New Review
        self.root.bind('<Control-n>', lambda e: self.new_review())
        self.root.bind('<Control-N>', lambda e: self.new_review())
        
        # Ctrl+B - Browse Reviews
        self.root.bind('<Control-b>', lambda e: self.browse_reviews())
        self.root.bind('<Control-B>', lambda e: self.browse_reviews())
        
        # Ctrl+Shift+A - Analytics (Shift to avoid Select All conflict)
        self.root.bind('<Control-Shift-A>', lambda e: self.show_analytics_dashboard())
        self.root.bind('<Control-Shift-a>', lambda e: self.show_analytics_dashboard())
        
        # Ctrl+D - Development Plan
        self.root.bind('<Control-d>', lambda e: self.show_idp_wizard())
        self.root.bind('<Control-D>', lambda e: self.show_idp_wizard())
        
        # Ctrl+O - Open Review
        self.root.bind('<Control-o>', lambda e: self.open_review())
        self.root.bind('<Control-O>', lambda e: self.open_review())
        
        # F1 - About
        self.root.bind('<F1>', lambda e: self.show_about_dialog())
        
        # Ctrl+Q - Quit
        self.root.bind('<Control-q>', lambda e: self.root.quit())
        self.root.bind('<Control-Q>', lambda e: self.root.quit())
        
        # Enter key - activate focused button
        def activate_button(event):
            widget = self.root.focus_get()
            if isinstance(widget, tk.Button):
                widget.invoke()
                return "break"
        self.root.bind('<Return>', activate_button)
    
    def update_status(self, message):
        """Update the status bar message"""
        if hasattr(self, 'status_bar'):
            self.status_bar.config(text=message)
            self.root.update_idletasks()
    
    def status_ready(self):
        """Set status to ready"""
        self.update_status("Ready")
    
    def status_loading(self):
        """Set status to loading"""
        self.update_status("Loading...")
    
    def status_saving(self):
        """Set status to saving"""
        self.update_status("Saving...")
    
    def status_info(self, count=None):
        """Show review count in status"""
        if count is None:
            try:
                # Exclude idp.json from count
                json_files = [f for f in self.reviews_dir.glob("*.json") if f.name != "idp.json"]
                count = len(json_files)
            except:
                count = 0
        self.update_status(f"Ready  •  {count} reviews saved")
    
    def create_menu_bar(self):
        """Create the menu bar"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # File menu
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="Home", command=self.show_home_screen, accelerator="Ctrl+H")
        file_menu.add_separator()
        file_menu.add_command(label="New Review", command=self.new_review, accelerator="Ctrl+N")
        file_menu.add_command(label="Open Review...", command=self.open_review, accelerator="Ctrl+O")
        file_menu.add_command(label="Browse All Reviews...", command=self.browse_reviews, accelerator="Ctrl+B")
        file_menu.add_separator()
        file_menu.add_command(label="View Analytics", command=self.show_analytics_dashboard, accelerator="Ctrl+Shift+A")
        file_menu.add_separator()
        file_menu.add_command(label="My Development Plan", command=self.show_idp_wizard, accelerator="Ctrl+D")
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.quit, accelerator="Ctrl+Q")
        
        # Settings menu
        settings_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Settings", menu=settings_menu)
        settings_menu.add_command(label="Personal Info", command=self.show_personal_settings)
        
        # Help menu
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Help", menu=help_menu)
        help_menu.add_command(label="Report Issue", command=self.open_github_issues)
        help_menu.add_separator()
        help_menu.add_command(label="About", command=self.show_about_dialog, accelerator="F1")
    
    def new_review(self):
        """Start a new review"""
        if messagebox.askyesno("New Review", "Start a new review? Any unsaved data will be lost."):
            self.session = ReviewSession()
            self.show_metadata_entry()
    
    def open_review(self):
        """Open an existing review"""
        self.status_loading()
        filename = filedialog.askopenfilename(
            title="Open Review",
            initialdir=self.reviews_dir,
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        
        if filename and self.load_review_json(filename):
            self.clear_content()  # Explicitly clear before loading
            messagebox.showinfo("Success", "Review loaded successfully!")
            self.show_metadata_entry()
        
        self.status_ready()
    
    def browse_reviews(self):
        """Show review browser dialog"""
        browser = tk.Toplevel(self.root)
        browser.title("Browse Reviews")
        browser.geometry("800x600")
        browser.transient(self.root)
        
        # Center the dialog
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - 400
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - 300
        browser.geometry(f"800x600+{x}+{y}")
        
        # Header
        header = tk.Frame(browser, bg="#1976D2", height=60)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        tk.Label(header, text="Review History", 
                font=("Segoe UI", 16, "bold"), bg="#1976D2", fg="white").pack(pady=15)
        
        # Main content
        content = tk.Frame(browser, bg="white")
        content.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Search frame
        search_frame = tk.Frame(content, bg="white")
        search_frame.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(search_frame, text="Search:", font=("Segoe UI", 10), bg="white").pack(side=tk.LEFT, padx=(0, 5))
        search_var = tk.StringVar()
        search_entry = tk.Entry(search_frame, textvariable=search_var, font=("Segoe UI", 10), width=30)
        search_entry.pack(side=tk.LEFT, padx=(0, 10))
        
        def update_list(*args):
            populate_list(search_var.get())
        
        search_var.trace_add('write', update_list)
        
        tk.Button(search_frame, text="Refresh", command=lambda: populate_list(search_var.get()),
                 font=("Segoe UI", 9), bg="#E0E0E0", fg="#212121",
                 padx=10, pady=5, relief=tk.FLAT, cursor="hand2").pack(side=tk.LEFT)
        
        # Listbox with scrollbar
        list_frame = tk.Frame(content, bg="white")
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        listbox = tk.Listbox(list_frame, font=("Segoe UI", 10), yscrollcommand=scrollbar.set)
        listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=listbox.yview)
        
        # Store file paths
        review_files = []
        
        def populate_list(search_term=""):
            """Populate listbox with reviews"""
            nonlocal review_files
            listbox.delete(0, tk.END)
            review_files = []
            
            # Get all JSON files
            try:
                json_files = sorted(self.reviews_dir.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
            except:
                json_files = []
            
            for filepath in json_files:
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    
                    referee = data.get('metadata', {}).get('referee', 'Unknown')
                    date = data.get('metadata', {}).get('date', 'Unknown')
                    game = data.get('metadata', {}).get('game_grade', 'Unknown')
                    
                    # Apply search filter
                    if search_term:
                        search_lower = search_term.lower()
                        if not (search_lower in referee.lower() or 
                               search_lower in date.lower() or 
                               search_lower in game.lower()):
                            continue
                    
                    display_text = f"{date} - {referee} - {game}"
                    listbox.insert(tk.END, display_text)
                    review_files.append(filepath)
                except:
                    pass
            
            if listbox.size() == 0:
                listbox.insert(tk.END, "No reviews found")
        
        def load_selected():
            """Load selected review"""
            selection = listbox.curselection()
            if not selection:
                messagebox.showwarning("No Selection", "Please select a review to load")
                return
            
            idx = selection[0]
            if idx < len(review_files):
                if self.load_review_json(review_files[idx]):
                    browser.destroy()
                    self.clear_content()  # Explicitly clear before loading
                    messagebox.showinfo("Success", "Review loaded successfully!")
                    self.show_metadata_entry()
        
        def delete_selected():
            """Delete selected review"""
            selection = listbox.curselection()
            if not selection:
                messagebox.showwarning("No Selection", "Please select a review to delete")
                return
            
            idx = selection[0]
            if idx < len(review_files):
                if messagebox.askyesno("Confirm Delete", "Are you sure you want to delete this review?"):
                    try:
                        review_files[idx].unlink()
                        populate_list(search_var.get())
                        messagebox.showinfo("Deleted", "Review deleted successfully")
                    except Exception as e:
                        messagebox.showerror("Error", f"Failed to delete:\n{str(e)}")
        
        # Buttons
        button_frame = tk.Frame(content, bg="white")
        button_frame.pack(fill=tk.X, pady=(10, 0))
        
        tk.Button(button_frame, text="Load", command=load_selected,
                 font=("Segoe UI", 10, "bold"), bg="#1976D2", fg="white",
                 padx=20, pady=8, relief=tk.FLAT, cursor="hand2").pack(side=tk.LEFT, padx=(0, 5))
        
        tk.Button(button_frame, text="Delete", command=delete_selected,
                 font=("Segoe UI", 10), bg="#F44336", fg="white",
                 padx=20, pady=8, relief=tk.FLAT, cursor="hand2").pack(side=tk.LEFT, padx=5)
        
        tk.Button(button_frame, text="Close", command=browser.destroy,
                 font=("Segoe UI", 10), bg="#E0E0E0", fg="#212121",
                 padx=20, pady=8, relief=tk.FLAT, cursor="hand2").pack(side=tk.RIGHT)
        
        # Populate initially
        populate_list()
        
        # Double-click to load
        listbox.bind('<Double-Button-1>', lambda e: load_selected())
    
    def show_analytics_dashboard(self):
        """Show analytics dashboard with charts and statistics"""
        # Load all reviews
        reviews = []
        try:
            json_files = sorted(self.reviews_dir.glob("*.json"), key=lambda p: p.stat().st_mtime)
        except:
            json_files = []
        
        for filepath in json_files:
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    reviews.append(data)
            except:
                pass
        
        if len(reviews) == 0:
            messagebox.showinfo("No Data", "No reviews found. Complete some reviews first!")
            return
        
        # Create dashboard window
        dashboard = tk.Toplevel(self.root)
        dashboard.title("Analytics Dashboard")
        dashboard.geometry("1200x800")
        dashboard.transient(self.root)
        
        # Center the dialog
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - 600
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - 400
        dashboard.geometry(f"1200x800+{x}+{y}")
        
        # Header
        header = tk.Frame(dashboard, bg="#1976D2", height=60)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        tk.Label(header, text="📊 Performance Analytics", 
                font=("Segoe UI", 18, "bold"), bg="#1976D2", fg="white").pack(pady=15)
        
        # Create scrollable content
        canvas = tk.Canvas(dashboard, bg="white", highlightthickness=0)
        scrollbar = ttk.Scrollbar(dashboard, orient="vertical", command=canvas.yview)
        content = tk.Frame(canvas, bg="white")
        
        content.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=content, anchor="nw", width=1150)
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Statistics Section
        stats_frame = tk.Frame(content, bg="white")
        stats_frame.pack(fill=tk.X, padx=20, pady=20)
        
        tk.Label(stats_frame, text="Overview Statistics", 
                font=("Segoe UI", 16, "bold"), bg="white", fg="#1976D2").pack(anchor="w", pady=(0, 10))
        
        # Calculate stats
        total_reviews = len(reviews)
        dates = [r.get('metadata', {}).get('date', '') for r in reviews if r.get('metadata', {}).get('date')]
        date_range = f"{min(dates)} to {max(dates)}" if dates else "N/A"
        
        grades = [r.get('metadata', {}).get('game_grade', '') for r in reviews if r.get('metadata', {}).get('game_grade')]
        most_common = max(set(grades), key=grades.count) if grades else "N/A"
        
        avg_difficulty = sum([r.get('difficulty', 5) for r in reviews]) / len(reviews) if reviews else 0
        
        # Display stats in grid
        stats_grid = tk.Frame(stats_frame, bg="white")
        stats_grid.pack(fill=tk.X)
        
        stats_data = [
            ("Total Reviews", str(total_reviews)),
            ("Date Range", date_range),
            ("Most Common Grade", most_common),
            ("Avg Difficulty", f"{avg_difficulty:.1f}/10")
        ]
        
        for i, (label, value) in enumerate(stats_data):
            stat_box = tk.Frame(stats_grid, bg="#E3F2FD", relief=tk.RAISED, bd=1)
            stat_box.grid(row=0, column=i, padx=5, pady=5, sticky="nsew")
            stats_grid.columnconfigure(i, weight=1)
            
            tk.Label(stat_box, text=label, font=("Segoe UI", 9), 
                    bg="#E3F2FD", fg="#757575").pack(pady=(10, 5))
            tk.Label(stat_box, text=value, font=("Segoe UI", 14, "bold"), 
                    bg="#E3F2FD", fg="#1976D2").pack(pady=(0, 10))
        
        # GFA Trends Chart
        chart_frame = tk.Frame(content, bg="white")
        chart_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        tk.Label(chart_frame, text="GFA Score Trends", 
                font=("Segoe UI", 16, "bold"), bg="white", fg="#1976D2").pack(anchor="w", pady=(0, 10))
        
        # Calculate average GFA scores per review
        review_dates = []
        avg_scores = []
        
        for review in reviews:
            date = review.get('metadata', {}).get('date', '')
            if not date:
                continue
            
            gfa_scores = review.get('gfa_scores', {})
            if gfa_scores:
                scores = [v for v in gfa_scores.values() if isinstance(v, (int, float))]
                if scores:
                    review_dates.append(date)
                    avg_scores.append(sum(scores) / len(scores))
        
        if review_dates and avg_scores:
            # Create matplotlib figure
            fig = Figure(figsize=(10, 4), dpi=100, facecolor='white')
            ax = fig.add_subplot(111)
            
            ax.plot(range(len(avg_scores)), avg_scores, marker='o', linewidth=2, 
                   markersize=8, color='#1976D2', label='Average GFA Score')
            ax.axhline(y=3, color='#FF9800', linestyle='--', label='Satisfactory (3.0)', alpha=0.7)
            ax.axhline(y=4, color='#4CAF50', linestyle='--', label='Sound (4.0)', alpha=0.7)
            
            ax.set_xlabel('Review Number', fontsize=11)
            ax.set_ylabel('Average Score', fontsize=11)
            ax.set_ylim(1, 5)
            ax.grid(True, alpha=0.3)
            ax.legend(loc='lower right')
            
            # Embed in tkinter
            canvas_widget = FigureCanvasTkAgg(fig, chart_frame)
            canvas_widget.draw()
            canvas_widget.get_tk_widget().pack(fill=tk.BOTH, expand=True)
        else:
            tk.Label(chart_frame, text="Not enough GFA data to display chart", 
                    font=("Segoe UI", 11), bg="white", fg="#757575").pack(pady=20)
        
        # Recent Reviews Table
        table_frame = tk.Frame(content, bg="white")
        table_frame.pack(fill=tk.X, padx=20, pady=20)
        
        tk.Label(table_frame, text="Recent Reviews", 
                font=("Segoe UI", 16, "bold"), bg="white", fg="#1976D2").pack(anchor="w", pady=(0, 10))
        
        # Create table
        table = tk.Frame(table_frame, bg="white")
        table.pack(fill=tk.X)
        
        headers = ["Date", "Game Grade", "Difficulty", "Avg GFA Score"]
        for i, header in enumerate(headers):
            tk.Label(table, text=header, font=("Segoe UI", 10, "bold"), 
                    bg="#1976D2", fg="white", padx=10, pady=8).grid(row=0, column=i, sticky="ew")
        
        # Show last 10 reviews
        for idx, review in enumerate(reversed(reviews[-10:])):
            date = review.get('metadata', {}).get('date', 'N/A')
            grade = review.get('metadata', {}).get('game_grade', 'N/A')
            diff = review.get('difficulty', 'N/A')
            
            gfa_scores = review.get('gfa_scores', {})
            if gfa_scores:
                scores = [v for v in gfa_scores.values() if isinstance(v, (int, float))]
                avg_gfa = f"{sum(scores)/len(scores):.1f}" if scores else "N/A"
            else:
                avg_gfa = "N/A"
            
            bg_color = "#F5F5F5" if idx % 2 == 0 else "white"
            
            tk.Label(table, text=date, font=("Segoe UI", 9), 
                    bg=bg_color, padx=10, pady=6).grid(row=idx+1, column=0, sticky="ew")
            tk.Label(table, text=grade, font=("Segoe UI", 9), 
                    bg=bg_color, padx=10, pady=6).grid(row=idx+1, column=1, sticky="ew")
            tk.Label(table, text=str(diff), font=("Segoe UI", 9), 
                    bg=bg_color, padx=10, pady=6).grid(row=idx+1, column=2, sticky="ew")
            tk.Label(table, text=avg_gfa, font=("Segoe UI", 9), 
                    bg=bg_color, padx=10, pady=6).grid(row=idx+1, column=3, sticky="ew")
        
        # Close button
        button_frame = tk.Frame(content, bg="white")
        button_frame.pack(fill=tk.X, padx=20, pady=20)
        
        tk.Button(button_frame, text="Close", command=dashboard.destroy,
                 font=("Segoe UI", 10, "bold"), bg="#1976D2", fg="white",
                 padx=30, pady=10, relief=tk.FLAT, cursor="hand2").pack()
        
        # Pack canvas and scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Bind mousewheel
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
    
    def show_about_dialog(self):
        """Show the About dialog"""
        about_window = tk.Toplevel(self.root)
        about_window.title("About")
        about_window.geometry("500x400")
        about_window.resizable(False, False)
        
        # Center the dialog
        about_window.transient(self.root)
        about_window.grab_set()
        
        # Calculate position
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - 250
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - 200
        about_window.geometry(f"500x400+{x}+{y}")
        
        # Header with color
        header = tk.Frame(about_window, bg="#1976D2", height=80)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        tk.Label(header, text="Rugby Referee Review System", 
                font=("Segoe UI", 18, "bold"), bg="#1976D2", fg="white").pack(pady=20)
        
        # Content
        content = tk.Frame(about_window, bg="white")
        content.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        tk.Label(content, text=f"Version {__version__}", 
                font=("Segoe UI", 11), bg="white", fg="#757575").pack(pady=5)
        
        tk.Label(content, text="", bg="white").pack(pady=5)  # Spacer
        
        tk.Label(content, text="Copyright © 2025 Andrew Clarkson", 
                font=("Segoe UI", 11, "bold"), bg="white", fg="#212121").pack(pady=5)
        
        tk.Label(content, text="All Rights Reserved", 
                font=("Segoe UI", 9), bg="white", fg="#757575").pack()
        
        tk.Label(content, text="", bg="white").pack(pady=10)  # Spacer
        
        tk.Label(content, text="Implements the Community Rugby Referee\nDevelopment Framework (CRRDF)", 
                font=("Segoe UI", 10), bg="white", fg="#212121", justify=tk.CENTER).pack(pady=5)
        
        tk.Label(content, text=f"CRRDF Framework: {CRRDF_VERSION}", 
                font=("Segoe UI", 9), bg="white", fg="#757575").pack(pady=2)
        
        tk.Label(content, text="", bg="white").pack(pady=10)  # Spacer
        
        tk.Label(content, text="Licensed under MIT License", 
                font=("Segoe UI", 9), bg="white", fg="#757575").pack()
        
        # Close button
        tk.Button(content, text="Close", command=about_window.destroy,
                 font=("Segoe UI", 10, "bold"), bg="#1976D2", fg="white",
                 padx=30, pady=8, relief=tk.FLAT, cursor="hand2").pack(pady=20)
    
    def open_github_issues(self):
        """Open GitHub issues page in browser"""
        import webbrowser
        github_url = "https://github.com/agclarkson/ReviewApp/issues"
        webbrowser.open(github_url)
    
    def show_personal_settings(self):
        """Show personal settings dialog"""
        settings = tk.Toplevel(self.root)
        settings.title("Personal Settings")
        settings.geometry("550x400")
        settings.resizable(False, False)
        
        # Center the dialog
        settings.transient(self.root)
        settings.grab_set()
        
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - 275
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - 200
        settings.geometry(f"550x400+{x}+{y}")
        
        # Header
        header = tk.Frame(settings, bg="#1976D2", height=60)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        tk.Label(header, text="Personal Information", 
                font=("Segoe UI", 16, "bold"), bg="#1976D2", fg="white").pack(pady=15)
        
        # Content
        content = tk.Frame(settings, bg="white")
        content.pack(fill=tk.BOTH, expand=True, padx=40, pady=20)
        
        tk.Label(content, text="This information auto-fills your review forms", 
                font=("Segoe UI", 9), bg="white", fg="#757575").pack(pady=(0, 20))
        
        # Name field
        tk.Label(content, text="Your Full Name:", 
                font=("Segoe UI", 10, "bold"), bg="white", fg="#212121", anchor="w").pack(fill=tk.X, pady=(0, 5))
        
        name_entry = tk.Entry(content, font=("Segoe UI", 11), width=40, relief=tk.SOLID, bd=1)
        name_entry.pack(fill=tk.X, pady=(0, 15), ipady=5)
        name_entry.insert(0, self.config.get("user_name", ""))
        
        # Coach field
        tk.Label(content, text="Preferred Coach Name (Optional):", 
                font=("Segoe UI", 10, "bold"), bg="white", fg="#212121", anchor="w").pack(fill=tk.X, pady=(0, 5))
        
        coach_entry = tk.Entry(content, font=("Segoe UI", 11), width=40, relief=tk.SOLID, bd=1)
        coach_entry.pack(fill=tk.X, pady=(0, 15), ipady=5)
        coach_entry.insert(0, self.config.get("coach_name", ""))
        
        # Status label
        status_label = tk.Label(content, text="", font=("Segoe UI", 10), bg="white", height=2)
        status_label.pack(pady=10)
        
        def save_settings():
            """Save updated settings"""
            name = name_entry.get().strip()
            
            if not name:
                status_label.config(text="⚠ Name cannot be empty", fg="#F44336")
                return
            
            coach = coach_entry.get().strip()
            
            self.config["user_name"] = name
            self.config["coach_name"] = coach
            self.save_config()
            
            status_label.config(text="✓ Settings saved successfully!", fg="#4CAF50")
            self.root.after(1500, settings.destroy)
        
        # Buttons
        button_frame = tk.Frame(content, bg="white")
        button_frame.pack(pady=5)
        
        tk.Button(button_frame, text="Save Changes", command=save_settings,
                 font=("Segoe UI", 10, "bold"), bg="#4CAF50", fg="white",
                 padx=30, pady=10, relief=tk.FLAT, cursor="hand2",
                 activebackground="#45a049").pack(side=tk.LEFT, padx=5)
        
        tk.Button(button_frame, text="Cancel", command=settings.destroy,
                 font=("Segoe UI", 10), bg="#E0E0E0", fg="#212121",
                 padx=30, pady=10, relief=tk.FLAT, cursor="hand2",
                 activebackground="#d0d0d0").pack(side=tk.LEFT, padx=5)
    
    def show_home_screen(self):
        """Show the home screen with main actions"""
        # Clear content area
        self.clear_content()
        
        # Update status
        self.status_info()
        
        # Main container
        home = tk.Frame(self.content, bg="white")
        home.pack(fill=tk.BOTH, expand=True)
        
        # Welcome section
        welcome_frame = tk.Frame(home, bg="white")
        welcome_frame.pack(pady=40)
        
        # Show icon if we have it (or placeholder)
        icon_frame = tk.Frame(welcome_frame, bg="white")
        icon_frame.pack(pady=20)
        
        # Create a simple placeholder icon using canvas
        canvas = tk.Canvas(icon_frame, width=120, height=120, bg="white", highlightthickness=0)
        canvas.pack()
        
        # Draw circle
        canvas.create_oval(10, 10, 110, 110, fill="#1976D2", outline="#0D47A1", width=3)
        
        # Draw text
        canvas.create_text(60, 60, text="RR", font=("Segoe UI", 36, "bold"), fill="white")
        
        # Welcome text
        user_name = self.config.get("user_name", "")
        welcome_text = f"Welcome back, {user_name}!" if user_name else "Welcome to Rugby Referee Review System"
        tk.Label(welcome_frame, text=welcome_text, 
                font=("Segoe UI", 18, "bold"), bg="white", fg="#212121").pack(pady=10)
        
        tk.Label(welcome_frame, text="Track your development with the CRRDF Framework", 
                font=("Segoe UI", 11), bg="white", fg="#757575").pack()
        
        # Action buttons
        buttons_frame = tk.Frame(home, bg="white")
        buttons_frame.pack(pady=30)
        
        # New Review button
        new_btn = tk.Button(buttons_frame, text="📝 New Review", 
                           command=self.show_metadata_entry,
                           font=("Segoe UI", 14, "bold"), bg="#4CAF50", fg="white",
                           padx=60, pady=20, relief=tk.FLAT, cursor="hand2",
                           activebackground="#45a049")
        new_btn.pack(pady=10, fill=tk.X)
        
        # Browse Reviews button
        browse_btn = tk.Button(buttons_frame, text="📂 Browse Reviews", 
                              command=self.browse_reviews,
                              font=("Segoe UI", 14, "bold"), bg="#2196F3", fg="white",
                              padx=60, pady=20, relief=tk.FLAT, cursor="hand2",
                              activebackground="#1976D2")
        browse_btn.pack(pady=10, fill=tk.X)
        
        # View Analytics button
        analytics_btn = tk.Button(buttons_frame, text="📊 View Analytics", 
                                 command=self.show_analytics_dashboard,
                                 font=("Segoe UI", 14, "bold"), bg="#FF9800", fg="white",
                                 padx=60, pady=20, relief=tk.FLAT, cursor="hand2",
                                 activebackground="#F57C00")
        analytics_btn.pack(pady=10, fill=tk.X)
        
        # Development Plan button - show "Edit" if exists, "Create" if not
        idp_text = "📋 Edit Development Plan" if self.idp_data else "📋 Create Development Plan"
        idp_btn = tk.Button(buttons_frame, text=idp_text, 
                           command=self.show_idp_wizard,
                           font=("Segoe UI", 14, "bold"), bg="#9C27B0", fg="white",
                           padx=60, pady=20, relief=tk.FLAT, cursor="hand2",
                           activebackground="#7B1FA2")
        idp_btn.pack(pady=10, fill=tk.X)
        
        # Stats and recent reviews section
        info_frame = tk.Frame(home, bg="#F5F5F5")
        info_frame.pack(fill=tk.X, padx=40, pady=20)
        
        # Quick stats
        stats_frame = tk.Frame(info_frame, bg="#F5F5F5")
        stats_frame.pack(pady=15)
        
        try:
            # Count reviews (exclude idp.json)
            json_files = [f for f in self.reviews_dir.glob("*.json") if f.name != "idp.json"]
            total_reviews = len(json_files)
            
            if total_reviews > 0:
                # Get latest review date
                latest_file = max(json_files, key=lambda p: p.stat().st_mtime)
                try:
                    with open(latest_file, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        latest_date = data.get('metadata', {}).get('date', 'Unknown')
                except:
                    latest_date = 'Unknown'
                
                tk.Label(stats_frame, text=f"📈 {total_reviews} Reviews Completed", 
                        font=("Segoe UI", 11, "bold"), bg="#F5F5F5", fg="#212121").pack(side=tk.LEFT, padx=20)
                tk.Label(stats_frame, text=f"📅 Latest: {latest_date}", 
                        font=("Segoe UI", 11), bg="#F5F5F5", fg="#757575").pack(side=tk.LEFT, padx=20)
            else:
                tk.Label(stats_frame, text="🎯 Start your first review to track your progress!", 
                        font=("Segoe UI", 11), bg="#F5F5F5", fg="#757575").pack()
        except:
            tk.Label(stats_frame, text="Ready to start tracking your development", 
                    font=("Segoe UI", 11), bg="#F5F5F5", fg="#757575").pack()
        
        # Recent reviews list
        recent_frame = tk.Frame(info_frame, bg="#F5F5F5")
        recent_frame.pack(pady=10, padx=20, fill=tk.BOTH, expand=True)
        
        tk.Label(recent_frame, text="Recent Reviews:", 
                font=("Segoe UI", 11, "bold"), bg="#F5F5F5", fg="#212121", anchor="w").pack(fill=tk.X, pady=(5, 10))
        
        try:
            # Get recent reviews (exclude idp.json)
            all_files = [f for f in self.reviews_dir.glob("*.json") if f.name != "idp.json"]
            json_files = sorted(all_files, key=lambda p: p.stat().st_mtime, reverse=True)[:5]
            
            if json_files:
                for filepath in json_files:
                    try:
                        with open(filepath, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                            date = data.get('metadata', {}).get('date', 'Unknown')
                            grade = data.get('metadata', {}).get('game_grade', 'Unknown')
                            
                            review_frame = tk.Frame(recent_frame, bg="white", relief=tk.RAISED, bd=1)
                            review_frame.pack(fill=tk.X, pady=2)
                            
                            review_btn = tk.Button(review_frame, 
                                                  text=f"  • {date} - {grade}", 
                                                  font=("Segoe UI", 10), bg="white", fg="#212121",
                                                  anchor="w", relief=tk.FLAT, cursor="hand2",
                                                  command=lambda f=filepath: self.load_review_from_home(f))
                            review_btn.pack(fill=tk.X, padx=10, pady=5)
                    except:
                        pass
            else:
                tk.Label(recent_frame, text="No reviews yet. Click 'New Review' to get started!", 
                        font=("Segoe UI", 10), bg="#F5F5F5", fg="#757575", anchor="w").pack(fill=tk.X)
        except:
            tk.Label(recent_frame, text="No reviews found", 
                    font=("Segoe UI", 10), bg="#F5F5F5", fg="#757575", anchor="w").pack(fill=tk.X)
    
    def load_review_from_home(self, filepath):
        """Load a review from the home screen recent list"""
        if self.load_review_json(filepath):
            self.show_metadata_entry()
    
    def create_widgets(self):
        """Create the main UI"""
        # Header with modern color
        header = tk.Frame(self.root, bg="#1976D2", height=80)
        header.pack(fill=tk.X)
        header.pack_propagate(False)  # Maintain fixed height
        
        title = tk.Label(header, text="Rugby Referee Review System", 
                        font=("Segoe UI", 18, "bold"), bg="#1976D2", fg="white")
        title.pack(pady=(15, 5))
        
        # Subtitle with copyright
        subtitle = tk.Label(header, text="Based on CRRDF Framework  •  © 2025 Andrew Clarkson", 
                           font=("Segoe UI", 9), bg="#1976D2", fg="#E3F2FD")
        subtitle.pack()
        
        # Status bar at bottom
        self.status_bar = tk.Label(self.root, text="Ready", bd=1, relief=tk.SUNKEN, anchor=tk.W,
                                   font=("Segoe UI", 9), bg="#F5F5F5", fg="#757575", padx=10, pady=3)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # Main content area
        self.content = tk.Frame(self.root, bg="#FAFAFA")
        self.content.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Start with metadata entry
        self.show_metadata_entry()
        
    def show_metadata_entry(self):
        """Game metadata entry screen"""
        self.clear_content()
        
        # Create scrollable canvas
        canvas = tk.Canvas(self.content, bg="#FAFAFA", highlightthickness=0)
        scrollbar = ttk.Scrollbar(self.content, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg="#FAFAFA")
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw", width=1050)
        canvas.configure(yscrollcommand=scrollbar.set)
        
        frame = scrollable_frame
        
        tk.Label(frame, text="Game Information", font=("Segoe UI", 16, "bold"),
                bg="#FAFAFA", fg="#212121").pack(pady=10)
        
        # Input fields
        self.metadata_entries = {}
        
        # Game & Grade - Dropdown
        container = tk.Frame(frame, bg="#FAFAFA")
        container.pack(fill=tk.X, pady=5, padx=20)
        tk.Label(container, text="Game & Grade:", width=15, anchor="w",
                bg="#FAFAFA", font=("Segoe UI", 10, "bold"), fg="#212121").pack(side=tk.LEFT)
        
        grade_var = tk.StringVar(value=self.session.metadata.get('game_grade', ''))
        grade_combo = ttk.Combobox(container, textvariable=grade_var, values=GAME_GRADES, 
                                   font=("Segoe UI", 10), width=38, state='readonly')
        grade_combo.pack(side=tk.LEFT, padx=5)
        if not self.session.metadata.get('game_grade'):
            grade_combo.set('Select grade...')
        
        # Create a wrapper to match PlaceholderEntry interface
        class ComboWrapper:
            def __init__(self, combo, var):
                self.combo = combo
                self.var = var
            def get_value(self):
                val = self.var.get()
                return "" if val == "Select grade..." else val
        
        self.metadata_entries['game_grade'] = ComboWrapper(grade_combo, grade_var)
        
        # Date - Manual entry with format hint
        container = tk.Frame(frame, bg="#FAFAFA")
        container.pack(fill=tk.X, pady=5, padx=20)
        tk.Label(container, text="Date:", width=15, anchor="w",
                bg="#FAFAFA", font=("Segoe UI", 10, "bold"), fg="#212121").pack(side=tk.LEFT)
        
        # Parse existing date or use today
        date_value = self.session.metadata.get('date', datetime.now().strftime("%Y-%m-%d"))
        
        date_entry = PlaceholderEntry(container, placeholder="YYYY-MM-DD", width=40, font=("Segoe UI", 10))
        date_entry.pack(side=tk.LEFT, padx=5)
        
        # Set default value
        if date_value:
            date_entry.delete(0, tk.END)
            date_entry.insert(0, date_value)
            date_entry.config(fg=date_entry.default_fg_color)
        
        # Add "Today" button
        def set_today():
            date_entry.delete(0, tk.END)
            date_entry.insert(0, datetime.now().strftime("%Y-%m-%d"))
            date_entry.config(fg=date_entry.default_fg_color)
        
        tk.Button(container, text="Today", command=set_today,
                 font=("Segoe UI", 9), bg="#E0E0E0", fg="#212121",
                 padx=10, pady=4, relief=tk.FLAT, cursor="hand2").pack(side=tk.LEFT, padx=5)
        
        self.metadata_entries['date'] = date_entry
        
        # Rest of fields - Text entries
        other_fields = [
            ("Result:", "result", "e.g., Home 24-17 Away"),
            ("Referee:", "referee", "Your name"),
            ("Coach:", "coach", "Coach name (optional)")
        ]
        
        for label, key, placeholder in other_fields:
            container = tk.Frame(frame, bg="#FAFAFA")
            container.pack(fill=tk.X, pady=5, padx=20)
            
            tk.Label(container, text=label, width=15, anchor="w",
                    bg="#FAFAFA", font=("Segoe UI", 10, "bold"), fg="#212121").pack(side=tk.LEFT)
            
            entry = PlaceholderEntry(container, placeholder=placeholder, width=40, font=("Segoe UI", 10))
            entry.pack(side=tk.LEFT, padx=5)
            self.metadata_entries[key] = entry
            
            # Populate with loaded data if it exists
            if self.session.metadata.get(key):
                entry.delete(0, tk.END)
                entry.insert(0, self.session.metadata[key])
                entry.config(fg=entry.default_fg_color)
            # Otherwise auto-fill from config for new reviews
            elif key == "referee" and self.config.get("user_name"):
                entry.delete(0, tk.END)
                entry.insert(0, self.config.get("user_name"))
                entry.config(fg=entry.default_fg_color)
            elif key == "coach" and self.config.get("coach_name"):
                entry.delete(0, tk.END)
                entry.insert(0, self.config.get("coach_name"))
                entry.config(fg=entry.default_fg_color)
        
        # Goals section
        tk.Label(frame, text="\nMatch Goals", font=("Segoe UI", 14, "bold"),
                bg="#FAFAFA", fg="#212121").pack(pady=10)
        
        tk.Label(frame, text="Primary Goal:", anchor="w", bg="#FAFAFA",
                font=("Segoe UI", 10, "bold"), fg="#212121").pack(fill=tk.X, padx=20)
        self.primary_goal = PlaceholderText(frame, placeholder="What is your main focus for this game? e.g., 'Improve accuracy in jackler decisions at breakdown'", 
                                           height=2, width=60, font=("Segoe UI", 10), wrap=tk.WORD)
        self.primary_goal.pack(padx=20, pady=5)
        
        # Populate with loaded data
        if self.session.goals.get('primary'):
            self.primary_goal.delete("1.0", tk.END)
            self.primary_goal.insert("1.0", self.session.goals['primary'])
            self.primary_goal.config(fg=self.primary_goal.default_fg_color)
        
        tk.Label(frame, text="Secondary Goal:", anchor="w", bg="#FAFAFA",
                font=("Segoe UI", 10, "bold"), fg="#212121").pack(fill=tk.X, padx=20)
        self.secondary_goal = PlaceholderText(frame, placeholder="Your secondary focus area? e.g., 'Better positioning at scrum time'", 
                                             height=2, width=60, font=("Segoe UI", 10), wrap=tk.WORD)
        self.secondary_goal.pack(padx=20, pady=5)
        
        # Populate with loaded data
        if self.session.goals.get('secondary'):
            self.secondary_goal.delete("1.0", tk.END)
            self.secondary_goal.insert("1.0", self.session.goals['secondary'])
            self.secondary_goal.config(fg=self.secondary_goal.default_fg_color)
        
        # Difficulty scale
        tk.Label(frame, text="\nGame Difficulty (1=Easy, 10=Very Hard):",
                bg="#FAFAFA", font=("Segoe UI", 10, "bold"), fg="#212121").pack(pady=10)
        
        self.difficulty_var = tk.IntVar(value=self.session.difficulty)
        
        # Display current value
        difficulty_value_label = tk.Label(frame, text=f"Current: {self.session.difficulty}", 
                                         font=("Segoe UI", 10), bg="#FAFAFA", fg="#1976D2")
        difficulty_value_label.pack()
        
        def update_difficulty_label(val):
            difficulty_value_label.config(text=f"Current: {val}")
        
        difficulty_scale = tk.Scale(frame, from_=1, to=10, orient=tk.HORIZONTAL,
                                   variable=self.difficulty_var, length=300, 
                                   bg="#FAFAFA", highlightthickness=0,
                                   command=update_difficulty_label)
        difficulty_scale.pack()
        
        # Navigation with modern styling
        nav = tk.Frame(frame, bg="#FAFAFA")
        nav.pack(pady=20)
        
        tk.Button(nav, text="Next: Self Reflection", command=self.show_self_reflection,
                 font=("Segoe UI", 11, "bold"), bg="#00796B", fg="white",
                 padx=20, pady=10, relief=tk.FLAT, cursor="hand2").pack()
        
        # Pack canvas and scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Bind mousewheel
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
    def show_self_reflection(self):
        """Self reflection questions"""
        # Save metadata
        for key, entry in self.metadata_entries.items():
            self.session.metadata[key] = entry.get_value()
        self.session.goals["primary"] = self.primary_goal.get_value()
        self.session.goals["secondary"] = self.secondary_goal.get_value()
        self.session.difficulty = self.difficulty_var.get()
        
        self.clear_content()
        
        # Create scrollable canvas
        canvas = tk.Canvas(self.content, bg="#FAFAFA", highlightthickness=0)
        scrollbar = ttk.Scrollbar(self.content, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg="#FAFAFA")
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw", width=1050)
        canvas.configure(yscrollcommand=scrollbar.set)
        
        frame = scrollable_frame
        
        tk.Label(frame, text="Self Reflection", font=("Segoe UI", 16, "bold"),
                bg="#FAFAFA", fg="#212121").pack(pady=10)
        
        questions = [
            ("Did I meet my goals - why or why not? Give examples:", "goals_met", 
             "Think about specific moments... e.g., 'Primary goal met - identified 8/10 jackler situations correctly. Secondary goal partly met - positioning at scrums good in first half but...'"),
            ("What areas went well and why? (relate to GFAs if possible):", "what_went_well",
             "Be specific... e.g., 'Tackle area management strong - clear communication on tackler release. Called advantage well when Blue under pressure in 34th min...'"),
            ("What was the biggest challenge? What would/could I do differently?:", "biggest_challenge",
             "Describe the challenge and your solution... e.g., 'Breakdown became messy in final 20 mins. Next time I'll be more proactive with warnings when I see the trend starting...'"),
            ("What am I taking into my next game?:", "taking_forward",
             "Your action points... e.g., 'Work on scrum positioning - stay wider to see both props. Keep using advantage effectively as today showed it helps game flow...'")
        ]
        
        self.reflection_entries = {}
        
        for question, key, placeholder in questions:
            tk.Label(frame, text=question, anchor="w", bg="#FAFAFA",
                    font=("Segoe UI", 10, "bold"), fg="#212121").pack(fill=tk.X, padx=20, pady=(10, 5))
            
            text = PlaceholderText(frame, placeholder=placeholder, height=4, width=70, 
                                  font=("Segoe UI", 10), wrap=tk.WORD)
            text.pack(padx=20, pady=5)
            self.reflection_entries[key] = text
            
            # Populate with loaded data
            if self.session.reflections.get(key):
                text.delete("1.0", tk.END)
                text.insert("1.0", self.session.reflections[key])
                text.config(fg=text.default_fg_color)
        
        # Navigation
        nav = tk.Frame(frame, bg="#FAFAFA")
        nav.pack(pady=20)
        
        tk.Button(nav, text="← Back", command=self.show_metadata_entry,
                 font=("Segoe UI", 10), bg="#E0E0E0", fg="#212121",
                 padx=15, pady=8, relief=tk.FLAT, cursor="hand2").pack(side=tk.LEFT, padx=5)
        tk.Button(nav, text="Next: CRRDF Deep Dive →", command=self.start_crrdf_questions,
                 font=("Segoe UI", 11, "bold"), bg="#00796B", fg="white",
                 padx=20, pady=10, relief=tk.FLAT, cursor="hand2").pack(side=tk.LEFT, padx=5)
        
        # Pack canvas and scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Bind mousewheel
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
    
    def start_crrdf_questions(self):
        """Begin CRRDF pillar-by-pillar questions"""
        # Save reflections
        for key, text_widget in self.reflection_entries.items():
            self.session.reflections[key] = text_widget.get_value()
        
        # Start with technical pillar
        self.pillars = list(CRRDF_QUESTIONS.keys())
        self.current_pillar_index = 0
        self.show_pillar_intro()
    
    def show_pillar_intro(self):
        """Show introduction to current pillar"""
        if self.current_pillar_index >= len(self.pillars):
            # All pillars complete, move to GFA scoring
            self.show_gfa_scoring()
            return
        
        pillar_name = self.pillars[self.current_pillar_index]
        pillar_data = CRRDF_QUESTIONS[pillar_name]
        
        self.clear_content()
        
        frame = tk.Frame(self.content, bg="#FAFAFA")
        frame.pack(fill=tk.BOTH, expand=True)
        
        # Pillar title
        tk.Label(frame, text=f"CRRDF: {pillar_data['title']}", 
                font=("Segoe UI", 18, "bold"), bg="#FAFAFA", fg="#1976D2").pack(pady=20)
        
        tk.Label(frame, text=f"This section has {len(pillar_data['questions'])} question(s).",
                font=("Segoe UI", 11), bg="#FAFAFA", fg="#757575").pack(pady=10)
        
        tk.Label(frame, text="Each question includes prompts to help you reflect deeply.\nBe specific with examples and time stamps where possible.",
                font=("Segoe UI", 10), bg="#FAFAFA", fg="#757575", justify=tk.CENTER).pack(pady=10)
        
        # Navigation
        nav = tk.Frame(frame, bg="#FAFAFA")
        nav.pack(pady=30)
        
        if self.current_pillar_index > 0:
            tk.Button(nav, text="← Back", command=self.previous_pillar,
                     font=("Segoe UI", 10), bg="#E0E0E0", fg="#212121",
                     padx=15, pady=8, relief=tk.FLAT, cursor="hand2").pack(side=tk.LEFT, padx=5)
        tk.Button(nav, text="Start Questions →", 
                 command=lambda: self.show_pillar_question(pillar_name, 0),
                 font=("Segoe UI", 11, "bold"), bg="#00796B", fg="white",
                 padx=20, pady=10, relief=tk.FLAT, cursor="hand2").pack(side=tk.LEFT, padx=5)
    
    def show_pillar_question(self, pillar_name, question_index):
        """Show a specific question from a pillar"""
        pillar_data = CRRDF_QUESTIONS[pillar_name]
        
        if question_index >= len(pillar_data['questions']):
            # Move to next pillar
            self.current_pillar_index += 1
            self.show_pillar_intro()
            return
        
        question_data = pillar_data['questions'][question_index]
        
        self.clear_content()
        
        frame = tk.Frame(self.content, bg="#FAFAFA")
        frame.pack(fill=tk.BOTH, expand=True)
        
        # Progress indicator
        progress_text = f"{pillar_data['title']} - Question {question_index + 1} of {len(pillar_data['questions'])}"
        tk.Label(frame, text=progress_text, font=("Segoe UI", 10), 
                bg="#FAFAFA", fg="#757575").pack(pady=5)
        
        # Question
        tk.Label(frame, text=question_data['q'], font=("Segoe UI", 13, "bold"),
                bg="#FAFAFA", fg="#212121", wraplength=800, justify=tk.LEFT).pack(pady=15, padx=20)
        
        # Prompts
        tk.Label(frame, text="Think about:", font=("Segoe UI", 10, "bold"),
                bg="#FAFAFA", fg="#00796B").pack(anchor=tk.W, padx=20)
        
        for prompt in question_data['prompts']:
            prompt_frame = tk.Frame(frame, bg="#FAFAFA")
            prompt_frame.pack(anchor=tk.W, padx=40, pady=2)
            tk.Label(prompt_frame, text="•", font=("Segoe UI", 10),
                    bg="#FAFAFA", fg="#00796B").pack(side=tk.LEFT)
            tk.Label(prompt_frame, text=prompt, font=("Segoe UI", 10),
                    bg="#FAFAFA", fg="#757575").pack(side=tk.LEFT, padx=5)
        
        # Answer area
        tk.Label(frame, text="\nYour Response:", font=("Segoe UI", 10, "bold"),
                bg="#FAFAFA", fg="#212121").pack(anchor=tk.W, padx=20, pady=(15, 5))
        
        # Get or create text widget for this question
        key = f"{pillar_name}_{question_index}"
        if not hasattr(self, 'pillar_answers'):
            self.pillar_answers = {}
        
        text = PlaceholderText(frame, placeholder="Be specific with examples, include time stamps where relevant...",
                              height=6, width=80, font=("Segoe UI", 10), wrap=tk.WORD)
        text.pack(padx=20, pady=5)
        
        # Restore previous answer if it exists
        if key in self.pillar_answers:
            text.delete("1.0", tk.END)
            text.insert("1.0", self.pillar_answers[key])
        
        self.current_answer_widget = text
        self.current_answer_key = key
        
        # Navigation
        nav = tk.Frame(frame, bg="#FAFAFA")
        nav.pack(pady=20)
        
        if question_index > 0:
            tk.Button(nav, text="← Previous", 
                     command=lambda: [self.save_current_answer(), 
                                    self.show_pillar_question(pillar_name, question_index - 1)],
                     font=("Segoe UI", 10), bg="#E0E0E0", fg="#212121",
                     padx=15, pady=8, relief=tk.FLAT, cursor="hand2").pack(side=tk.LEFT, padx=5)
        tk.Button(nav, text="Next →" if question_index < len(pillar_data['questions']) - 1 else "Complete Pillar →",
                 command=lambda: [self.save_current_answer(), 
                                self.show_pillar_question(pillar_name, question_index + 1)],
                 font=("Segoe UI", 11, "bold"), bg="#00796B", fg="white",
                 padx=20, pady=10, relief=tk.FLAT, cursor="hand2").pack(side=tk.LEFT, padx=5)
    
    def save_current_answer(self):
        """Save the current answer"""
        if hasattr(self, 'current_answer_widget') and hasattr(self, 'current_answer_key'):
            answer = self.current_answer_widget.get_value()
            self.pillar_answers[self.current_answer_key] = answer
            
            # Also save to session
            if not hasattr(self.session, 'crrdf_reflections'):
                self.session.crrdf_reflections = {}
            self.session.crrdf_reflections[self.current_answer_key] = answer
    
    def previous_pillar(self):
        """Go back to previous pillar"""
        if self.current_pillar_index > 0:
            self.current_pillar_index -= 1
            self.show_pillar_intro()
    
    def show_gfa_scoring(self):
        """GFA scoring interface"""
        self.clear_content()
        
        # Create scrollable frame
        canvas = tk.Canvas(self.content, bg="#FAFAFA", highlightthickness=0)
        scrollbar = ttk.Scrollbar(self.content, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg="#FAFAFA")
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Title
        tk.Label(scrollable_frame, text="Game Focus Area Scoring", 
                font=("Segoe UI", 18, "bold"), bg="#FAFAFA", fg="#1976D2").pack(pady=15)
        
        tk.Label(scrollable_frame, text="Rate your performance in each focus area (1-5 scale)",
                font=("Segoe UI", 11), bg="#FAFAFA", fg="#757575").pack(pady=5)
        
        # Rating scale reference
        scale_frame = tk.Frame(scrollable_frame, bg="#E3F2FD", relief=tk.RIDGE, bd=1)
        scale_frame.pack(pady=15, padx=20, fill=tk.X)
        
        tk.Label(scale_frame, text="Rating Scale Reference:", font=("Segoe UI", 10, "bold"),
                bg="#E3F2FD", fg="#1976D2").pack(pady=5)
        
        for score, description in RATING_SCALE.items():
            tk.Label(scale_frame, text=f"{score} - {description}", 
                    font=("Segoe UI", 9), bg="#E3F2FD", fg="#212121",
                    wraplength=700, justify=tk.LEFT).pack(anchor=tk.W, padx=10, pady=2)
        
        # GFA Categories
        self.gfa_vars = {}
        
        for category, aspects in GFA_CATEGORIES.items():
            # Category header
            cat_frame = tk.Frame(scrollable_frame, bg="#FFFFFF", relief=tk.RIDGE, bd=1)
            cat_frame.pack(pady=10, padx=20, fill=tk.X)
            
            tk.Label(cat_frame, text=category, font=("Segoe UI", 12, "bold"),
                    bg="#00796B", fg="white").pack(fill=tk.X, pady=5)
            
            for aspect_name, description in aspects:
                aspect_frame = tk.Frame(cat_frame, bg="#FFFFFF")
                aspect_frame.pack(fill=tk.X, padx=15, pady=8)
                
                # Aspect name and description
                label_frame = tk.Frame(aspect_frame, bg="#FFFFFF")
                label_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
                
                tk.Label(label_frame, text=aspect_name, font=("Segoe UI", 10, "bold"),
                        bg="#FFFFFF", fg="#212121", anchor="w").pack(fill=tk.X)
                tk.Label(label_frame, text=description, font=("Segoe UI", 9),
                        bg="#FFFFFF", fg="#757575", wraplength=500, 
                        justify=tk.LEFT, anchor="w").pack(fill=tk.X)
                
                # Rating scale
                rating_frame = tk.Frame(aspect_frame, bg="#FFFFFF")
                rating_frame.pack(side=tk.RIGHT)
                
                var = tk.IntVar(value=3)
                self.gfa_vars[f"{category}_{aspect_name}"] = var
                
                for i in range(1, 6):
                    tk.Radiobutton(rating_frame, text=str(i), variable=var, value=i,
                                  bg="#FFFFFF", font=("Segoe UI", 9),
                                  selectcolor="#4CAF50").pack(side=tk.LEFT, padx=3)
        
        # Navigation
        nav = tk.Frame(scrollable_frame, bg="#FAFAFA")
        nav.pack(pady=20)
        
        tk.Button(nav, text="← Back to CRRDF", command=self.back_to_last_pillar,
                 font=("Segoe UI", 10), bg="#E0E0E0", fg="#212121",
                 padx=15, pady=8, relief=tk.FLAT, cursor="hand2").pack(side=tk.LEFT, padx=5)
        tk.Button(nav, text="Save & Export", command=self.export_to_excel,
                 font=("Segoe UI", 11, "bold"), bg="#4CAF50", fg="white",
                 padx=20, pady=10, relief=tk.FLAT, cursor="hand2").pack(side=tk.LEFT, padx=5)
        
        # Pack canvas and scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Bind mousewheel
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
    
    def back_to_last_pillar(self):
        """Return to the last pillar"""
        self.current_pillar_index = len(self.pillars) - 1
        pillar_name = self.pillars[self.current_pillar_index]
        last_q = len(CRRDF_QUESTIONS[pillar_name]['questions']) - 1
        self.show_pillar_question(pillar_name, last_q)
    
    def export_to_excel(self):
        """Export review to Excel file"""
        # Save GFA scores
        for key, var in self.gfa_vars.items():
            self.session.gfa_scores[key] = var.get()
        
        # Auto-save JSON first
        json_filename = self.save_review_json()
        
        # Ask for Excel save location
        filename = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
            initialfile=f"Review_{self.session.metadata['referee']}_{self.session.metadata['date']}.xlsx"
        )
        
        if not filename:
            # Still saved JSON even if they cancel Excel
            if json_filename:
                messagebox.showinfo("Saved", f"Review saved to:\n{json_filename}")
            return
        
        try:
            self._create_excel_export(filename)
            if json_filename:
                messagebox.showinfo("Success", f"Review exported successfully!\n\nExcel: {filename}\nJSON: {json_filename}")
            else:
                messagebox.showinfo("Success", f"Review exported successfully to:\n{filename}")
            
            # Return to home screen after successful save
            self.show_home_screen()
            
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export:\n{str(e)}")
    
    def _create_excel_export(self, filename):
        """Create the Excel workbook"""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Referee Review"
        
        # Styles
        header_font = Font(bold=True, size=14, color="FFFFFF")
        header_fill = PatternFill(start_color="1976D2", end_color="1976D2", fill_type="solid")
        subheader_font = Font(bold=True, size=12)
        subheader_fill = PatternFill(start_color="E3F2FD", end_color="E3F2FD", fill_type="solid")
        
        border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        
        row = 1
        
        # Title
        ws.merge_cells(f'A{row}:G{row}')
        cell = ws.cell(row, 1, "Rugby Referee Review")
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
        row += 1
        
        # Copyright
        ws.merge_cells(f'A{row}:G{row}')
        cell = ws.cell(row, 1, f"Based on CRRDF Framework  •  © 2025 Andrew Clarkson")
        cell.font = Font(size=9, italic=True)
        cell.alignment = Alignment(horizontal='center')
        row += 2
        
        # Metadata
        ws.cell(row, 1, "Game Information").font = subheader_font
        ws.cell(row, 1).fill = subheader_fill
        row += 1
        
        for key, label in [("game_grade", "Game & Grade:"), ("date", "Date:"), 
                           ("result", "Result:"), ("referee", "Referee:"), ("coach", "Coach:")]:
            ws.cell(row, 1, label).font = Font(bold=True)
            ws.cell(row, 2, self.session.metadata.get(key, ""))
            row += 1
        
        row += 1
        ws.cell(row, 1, "Match Goals").font = subheader_font
        ws.cell(row, 1).fill = subheader_fill
        row += 1
        
        ws.cell(row, 1, "Primary:").font = Font(bold=True)
        ws.cell(row, 2, self.session.goals.get("primary", ""))
        row += 1
        
        ws.cell(row, 1, "Secondary:").font = Font(bold=True)
        ws.cell(row, 2, self.session.goals.get("secondary", ""))
        row += 1
        
        ws.cell(row, 1, "Difficulty (1-10):").font = Font(bold=True)
        ws.cell(row, 2, str(self.session.difficulty))
        row += 2
        
        # Self Reflections
        ws.cell(row, 1, "Self Reflection").font = subheader_font
        ws.cell(row, 1).fill = subheader_fill
        row += 1
        
        for label, key in [("Goals Met:", "goals_met"), ("What Went Well:", "what_went_well"),
                          ("Biggest Challenge:", "biggest_challenge"), ("Taking Forward:", "taking_forward")]:
            ws.cell(row, 1, label).font = Font(bold=True)
            ws.merge_cells(f'B{row}:G{row}')
            ws.cell(row, 2, self.session.reflections.get(key, ""))
            ws.cell(row, 2).alignment = Alignment(wrap_text=True, vertical='top')
            row += 1
        
        row += 1
        
        # CRRDF Reflections
        ws.cell(row, 1, "CRRDF Framework Responses").font = subheader_font
        ws.cell(row, 1).fill = subheader_fill
        row += 1
        
        if hasattr(self, 'pillar_answers'):
            for pillar_name in self.pillars:
                pillar_data = CRRDF_QUESTIONS[pillar_name]
                ws.cell(row, 1, pillar_data['title']).font = Font(bold=True, size=11)
                row += 1
                
                for q_idx, question_data in enumerate(pillar_data['questions']):
                    key = f"{pillar_name}_{q_idx}"
                    ws.cell(row, 1, f"Q: {question_data['q']}").font = Font(bold=True)
                    row += 1
                    
                    ws.merge_cells(f'B{row}:G{row}')
                    ws.cell(row, 2, self.pillar_answers.get(key, ""))
                    ws.cell(row, 2).alignment = Alignment(wrap_text=True, vertical='top')
                    row += 1
                
                row += 1
        
        # GFA Scores
        ws.cell(row, 1, "Game Focus Area Scores").font = subheader_font
        ws.cell(row, 1).fill = subheader_fill
        row += 1
        
        # GFA Headers
        ws.cell(row, 1, "Category").font = Font(bold=True)
        ws.cell(row, 2, "Aspect").font = Font(bold=True)
        ws.cell(row, 3, "What this covers").font = Font(bold=True)
        ws.cell(row, 4, "Self Score").font = Font(bold=True)
        ws.cell(row, 5, "Coach Score").font = Font(bold=True)
        row += 1
        
        for category, aspects in GFA_CATEGORIES.items():
            for aspect_name, description in aspects:
                ws.cell(row, 1, category)
                ws.cell(row, 2, aspect_name)
                ws.cell(row, 3, description)
                key = f"{category}_{aspect_name}"
                ws.cell(row, 4, self.session.gfa_scores.get(key, ""))
                row += 1
        
        row += 1
        
        # Rating Scale
        ws.cell(row, 1, "Rating Scale").font = subheader_font
        ws.cell(row, 1).fill = subheader_fill
        row += 1
        
        for score, description in RATING_SCALE.items():
            ws.cell(row, 1, str(score)).font = Font(bold=True)
            ws.merge_cells(f'B{row}:G{row}')
            ws.cell(row, 2, description)
            ws.cell(row, 2).alignment = Alignment(wrap_text=True)
            row += 1
        
        # Column widths
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 50
        ws.column_dimensions['C'].width = 60
        ws.column_dimensions['D'].width = 15
        ws.column_dimensions['E'].width = 15
        
        wb.save(filename)
    
    def clear_content(self):
        """Clear the content area"""
        for widget in self.content.winfo_children():
            widget.destroy()
    
    def show_idp_wizard(self):
        """Show IDP creation/editing wizard"""
        self.clear_content()
        self.update_status("Development Plan")
        
        # Initialize IDP if not exists
        if not self.idp_data:
            self.idp_data = IDPData()
            self.idp_data.referee_name = self.config.get("user_name", "")
        
        # Section names and questions
        sections = {
            "club": {"title": "Club Level", "desc": "Your current club rugby status"},
            "rep": {"title": "Representative Level", "desc": "Your rep rugby involvement"},
            "aspirations": {"title": "Aspirations", "desc": "Your goals and targets"},
            "fitness": {"title": "Fitness Plan", "desc": "Physical preparation"},
            "law": {"title": "Law Knowledge", "desc": "Laws and application"},
            "mental": {"title": "Mental Game", "desc": "Mental preparation and resilience"},
            "focus1": {"title": "Focus Area 1", "desc": "First key development goal"},
            "focus2": {"title": "Focus Area 2", "desc": "Second key development goal"},
            "focus3": {"title": "Focus Area 3", "desc": "Third key development goal"}
        }
        
        section_key = self.idp_sections[self.idp_current_section]
        section_info = sections[section_key]
        
        # Main container
        container = tk.Frame(self.content, bg="white")
        container.pack(fill=tk.BOTH, expand=True, padx=40, pady=20)
        
        # Progress indicator
        progress_frame = tk.Frame(container, bg="white")
        progress_frame.pack(fill=tk.X, pady=(0, 20))
        
        progress_text = f"Section {self.idp_current_section + 1} of {len(self.idp_sections)}"
        tk.Label(progress_frame, text=progress_text,
                font=("Segoe UI", 10), bg="white", fg="#757575").pack()
        
        # Progress bar
        progress_bar = tk.Frame(progress_frame, bg="#E0E0E0", height=6)
        progress_bar.pack(fill=tk.X, pady=5)
        
        percent = (self.idp_current_section / len(self.idp_sections)) * 100
        fill_width = int(percent)
        if fill_width > 0:
            tk.Frame(progress_bar, bg="#1976D2", width=fill_width, height=6).place(x=0, y=0)
        
        # Section title
        tk.Label(container, text=section_info["title"],
                font=("Segoe UI", 20, "bold"), bg="white", fg="#212121").pack(pady=(0, 5))
        tk.Label(container, text=section_info["desc"],
                font=("Segoe UI", 11), bg="white", fg="#757575").pack(pady=(0, 20))
        
        # Scrollable content
        scroll_container = tk.Frame(container, bg="white")
        scroll_container.pack(fill=tk.BOTH, expand=True)
        
        canvas = tk.Canvas(scroll_container, bg="white", highlightthickness=0)
        scrollbar = tk.Scrollbar(scroll_container, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg="white")
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Enable mousewheel scrolling
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        # Questions for this section
        self.show_idp_section_questions(scrollable_frame, section_key)
        
        # Navigation buttons
        nav_frame = tk.Frame(container, bg="white")
        nav_frame.pack(fill=tk.X, pady=(20, 0))
        
        if self.idp_current_section > 0:
            tk.Button(nav_frame, text="← Previous",
                     command=lambda: self.idp_navigate(-1),
                     font=("Segoe UI", 11), bg="#E0E0E0", fg="#212121",
                     padx=20, pady=10, relief=tk.FLAT, cursor="hand2").pack(side=tk.LEFT)
        
        tk.Button(nav_frame, text="Save Draft",
                 command=self.save_idp,
                 font=("Segoe UI", 11), bg="#2196F3", fg="white",
                 padx=20, pady=10, relief=tk.FLAT, cursor="hand2").pack(side=tk.LEFT, padx=10)
        
        if self.idp_current_section < len(self.idp_sections) - 1:
            tk.Button(nav_frame, text="Next →",
                     command=lambda: self.idp_navigate(1),
                     font=("Segoe UI", 11, "bold"), bg="#4CAF50", fg="white",
                     padx=20, pady=10, relief=tk.FLAT, cursor="hand2").pack(side=tk.RIGHT)
        else:
            tk.Button(nav_frame, text="Complete & Export",
                     command=self.idp_complete,
                     font=("Segoe UI", 11, "bold"), bg="#4CAF50", fg="white",
                     padx=20, pady=10, relief=tk.FLAT, cursor="hand2").pack(side=tk.RIGHT)
    
    def show_idp_section_questions(self, parent, section_key):
        """Show questions for specific IDP section"""
        # This will contain the guided questions for each section
        # For now, adding a simplified version - we'll expand this
        
        questions = self.get_idp_questions(section_key)
        self.idp_entries = {}
        
        for q_id, question in questions.items():
            # Question label
            q_frame = tk.Frame(parent, bg="white")
            q_frame.pack(fill=tk.X, pady=10)
            
            tk.Label(q_frame, text=question["text"],
                    font=("Segoe UI", 11, "bold"), bg="white", fg="#212121",
                    wraplength=600, justify="left", anchor="w").pack(fill=tk.X, pady=(0, 5))
            
            if question.get("help"):
                tk.Label(q_frame, text=question["help"],
                        font=("Segoe UI", 9), bg="white", fg="#757575",
                        wraplength=600, justify="left", anchor="w").pack(fill=tk.X, pady=(0, 5))
            
            # Input field
            if question.get("type") == "text":
                entry = tk.Text(q_frame, height=4, font=("Segoe UI", 10),
                               wrap=tk.WORD, relief=tk.SOLID, bd=1)
                entry.pack(fill=tk.X, pady=5)
                
                # Add tab navigation - move to next widget instead of inserting tab
                def focus_next(event):
                    event.widget.tk_focusNext().focus()
                    return "break"  # Prevent default tab behavior
                entry.bind("<Tab>", focus_next)
                
                # Load existing value
                field_name = f"{section_key}_{q_id}"
                value = getattr(self.idp_data, field_name, "")
                if value:
                    entry.insert("1.0", value)
                
                self.idp_entries[field_name] = entry
            
            elif question.get("type") == "dropdown":
                var = tk.StringVar()
                combo = ttk.Combobox(q_frame, textvariable=var,
                                    values=question["options"],
                                    state="readonly", font=("Segoe UI", 10))
                combo.pack(fill=tk.X, pady=5)
                
                # Load existing value
                field_name = f"{section_key}_{q_id}"
                value = getattr(self.idp_data, field_name, "")
                if value:
                    var.set(value)
                
                self.idp_entries[field_name] = var
    
    def get_idp_questions(self, section_key):
        """Get questions for a specific IDP section"""
        questions = {
            "club": {
                "level": {
                    "text": "What level do you currently referee at club rugby?",
                    "type": "text",
                    "help": "e.g., Division 1, U21 Colts, etc."
                },
                "description": {
                    "text": "How would you describe your current club performances?",
                    "type": "text",
                    "help": "What's going well? What challenges are you facing?"
                },
                "challenges": {
                    "text": "What are your main challenges at club level?",
                    "type": "text",
                    "help": "Be specific about what you find difficult"
                },
                "goal": {
                    "text": "Where do you want to be by end of season?",
                    "type": "text",
                    "help": "Be specific - which grade? What role?"
                }
            },
            "rep": {
                "involvement": {
                    "text": "What's your current representative involvement?",
                    "type": "text",
                    "help": "e.g., Assistant referee, touch judge, not involved yet"
                },
                "details": {
                    "text": "Tell us about your rep experience:",
                    "type": "text",
                    "help": "What competitions? What's going well? What would you like to do more of?"
                },
                "goals": {
                    "text": "What are your representative goals this year?",
                    "type": "text",
                    "help": "What role do you want? What competitions?"
                }
            },
            "aspirations": {
                "ultimate_goal": {
                    "text": "What's your ultimate refereeing goal?",
                    "type": "text",
                    "help": "Dream big - where do you ultimately want to referee?"
                },
                "season_goal": {
                    "text": "What's realistic for this season?",
                    "type": "text",
                    "help": "Break your big goal into yearly steps"
                },
                "success_criteria": {
                    "text": "What would success look like this year?",
                    "type": "text",
                    "help": "How will you know you've progressed?"
                }
            },
            "fitness": {
                "level": {
                    "text": "Can you keep up with play for 80 minutes?",
                    "type": "text",
                    "help": "Be honest about your current fitness level"
                },
                "bronco_done": {
                    "text": "Have you done a Bronco test?",
                    "type": "text",
                    "help": "Yes/No - if yes, what was your time?"
                },
                "bronco_time": {
                    "text": "Current Bronco time (if applicable):",
                    "type": "text",
                    "help": "e.g., 6:00"
                },
                "bronco_target": {
                    "text": "Target Bronco time:",
                    "type": "text",
                    "help": "What time are you aiming for? e.g., 5:30"
                },
                "training_frequency": {
                    "text": "How often will you train?",
                    "type": "text",
                    "help": "Be realistic - days per week"
                },
                "training_types": {
                    "text": "What training will you do?",
                    "type": "text",
                    "help": "e.g., Running, gym, sport (squash, etc.)"
                },
                "fitness_target": {
                    "text": "Your specific fitness target:",
                    "type": "text",
                    "help": "What do you want to achieve? Be specific and measurable"
                },
                "fitness_obstacles": {
                    "text": "What might stop you training? How will you overcome it?",
                    "type": "text",
                    "help": "Think about obstacles and solutions now"
                }
            },
            "law": {
                "confidence": {
                    "text": "How confident are you with the laws? (1-5)",
                    "type": "text",
                    "help": "1 = Need lots of work, 5 = Very confident"
                },
                "weak_areas": {
                    "text": "Which areas need work?",
                    "type": "text",
                    "help": "e.g., Scrum, breakdown, maul, offside, advantage"
                },
                "consistency": {
                    "text": "Can you apply laws consistently for 80 minutes?",
                    "type": "text",
                    "help": "What affects your consistency? Fatigue? Pressure?"
                },
                "plan": {
                    "text": "Your law development plan:",
                    "type": "text",
                    "help": "How often will you study? What resources? Who can help?"
                }
            },
            "mental": {
                "strengths": {
                    "text": "What are your mental strengths?",
                    "type": "text",
                    "help": "e.g., Staying calm, confidence, focus, decision-making"
                },
                "challenges": {
                    "text": "What affects your performance mentally?",
                    "type": "text",
                    "help": "e.g., Nerves, pressure, mistakes, external factors"
                },
                "plan": {
                    "text": "How will you build mental resilience?",
                    "type": "text",
                    "help": "What will you work on? Who can support you?"
                }
            },
            "focus1": {
                "category": {
                    "text": "Focus Area Category:",
                    "type": "text",
                    "help": "e.g., Fitness, Law, Mental, Positioning, Communication"
                },
                "area": {
                    "text": "Specific focus area:",
                    "type": "text",
                    "help": "What exactly will you focus on?"
                },
                "why": {
                    "text": "Why is this important to you?",
                    "type": "text",
                    "help": "How will this help you achieve your aspirations?"
                },
                "goal": {
                    "text": "Your SPECIFIC, MEASURABLE goal:",
                    "type": "text",
                    "help": "e.g., 'Achieve 5:30 Bronco by June' not 'Get fitter'"
                },
                "how": {
                    "text": "How will you achieve this?",
                    "type": "text",
                    "help": "Specific actions, frequency, timeline"
                },
                "track": {
                    "text": "How will you track progress?",
                    "type": "text",
                    "help": "How will you know you're improving?"
                },
                "obstacles": {
                    "text": "Obstacles and how you'll overcome them:",
                    "type": "text",
                    "help": "What might stop you? What's your plan B?"
                }
            },
            "focus2": {
                "category": {
                    "text": "Focus Area Category:",
                    "type": "text",
                    "help": "e.g., Fitness, Law, Mental, Positioning, Communication"
                },
                "area": {
                    "text": "Specific focus area:",
                    "type": "text",
                    "help": "What exactly will you focus on?"
                },
                "why": {
                    "text": "Why is this important to you?",
                    "type": "text",
                    "help": "How will this help you achieve your aspirations?"
                },
                "goal": {
                    "text": "Your SPECIFIC, MEASURABLE goal:",
                    "type": "text",
                    "help": "Make it something you can clearly succeed or fail at"
                },
                "how": {
                    "text": "How will you achieve this?",
                    "type": "text",
                    "help": "Specific actions, frequency, timeline"
                },
                "track": {
                    "text": "How will you track progress?",
                    "type": "text",
                    "help": "How will you know you're improving?"
                },
                "obstacles": {
                    "text": "Obstacles and how you'll overcome them:",
                    "type": "text",
                    "help": "What might stop you? What's your plan B?"
                }
            },
            "focus3": {
                "category": {
                    "text": "Focus Area Category:",
                    "type": "text",
                    "help": "e.g., Fitness, Law, Mental, Positioning, Communication"
                },
                "area": {
                    "text": "Specific focus area:",
                    "type": "text",
                    "help": "What exactly will you focus on?"
                },
                "why": {
                    "text": "Why is this important to you?",
                    "type": "text",
                    "help": "How will this help you achieve your aspirations?"
                },
                "goal": {
                    "text": "Your SPECIFIC, MEASURABLE goal:",
                    "type": "text",
                    "help": "Make it something you can clearly succeed or fail at"
                },
                "how": {
                    "text": "How will you achieve this?",
                    "type": "text",
                    "help": "Specific actions, frequency, timeline"
                },
                "track": {
                    "text": "How will you track progress?",
                    "type": "text",
                    "help": "How will you know you're improving?"
                },
                "obstacles": {
                    "text": "Obstacles and how you'll overcome them:",
                    "type": "text",
                    "help": "What might stop you? What's your plan B?"
                }
            }
        }
        
        return questions.get(section_key, {})
    
    def idp_navigate(self, direction):
        """Navigate between IDP sections"""
        # Save current section
        self.idp_save_current_section()
        
        # Move to next/previous section
        self.idp_current_section += direction
        self.idp_current_section = max(0, min(self.idp_current_section, len(self.idp_sections) - 1))
        
        # Show new section
        self.show_idp_wizard()
    
    def idp_save_current_section(self):
        """Save current section data to IDP"""
        if not hasattr(self, 'idp_entries'):
            return
        
        for field_name, entry in self.idp_entries.items():
            if isinstance(entry, tk.Text):
                value = entry.get("1.0", tk.END).strip()
            elif isinstance(entry, tk.StringVar):
                value = entry.get()
            else:
                value = ""
            
            if hasattr(self.idp_data, field_name):
                setattr(self.idp_data, field_name, value)
    
    def idp_complete(self):
        """Complete IDP and export"""
        # Save final section
        self.idp_save_current_section()
        
        # Save to file
        if self.save_idp():
            messagebox.showinfo("Success", "IDP saved successfully!")
            
            # Offer to export
            if messagebox.askyesno("Export", "Would you like to export your IDP to Word?"):
                self.idp_export_word()
            
            # Return to home
            self.show_home_screen()
    
    def idp_export_word(self):
        """Export IDP to Word document matching ORRA template"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            messagebox.showerror("Error", "python-docx not installed.\nRun: pip install python-docx")
            return
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            initialfile=f"IDP_{self.idp_data.referee_name.replace(' ', '_')}_{self.idp_data.date}.docx"
        )
        
        if not filename:
            return
        
        try:
            doc = Document()
            
            # Title
            title = doc.add_paragraph()
            title_run = title.add_run("Referee Individual Development Plan")
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Referee details
            doc.add_paragraph(f"Referee: {self.idp_data.referee_name}")
            doc.add_paragraph(f"Date: {self.idp_data.date}")
            doc.add_paragraph()
            
            # Table 1: Current Reality
            table1 = doc.add_table(rows=2, cols=3)
            table1.style = 'Table Grid'
            
            # Headers
            headers = ["Club", "Representative", "Aspirations"]
            for i, header in enumerate(headers):
                cell = table1.rows[0].cells[i]
                cell.text = header
                # Bold the header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Content - combine all relevant fields for each column
            club_content = []
            if self.idp_data.club_level:
                club_content.append(self.idp_data.club_level)
            if self.idp_data.club_description:
                club_content.append(self.idp_data.club_description)
            if self.idp_data.club_challenges:
                club_content.append(f"Challenges: {self.idp_data.club_challenges}")
            if self.idp_data.club_goal:
                club_content.append(f"Goal: {self.idp_data.club_goal}")
            
            rep_content = []
            if self.idp_data.rep_involvement:
                rep_content.append(self.idp_data.rep_involvement)
            if self.idp_data.rep_details:
                rep_content.append(self.idp_data.rep_details)
            if self.idp_data.rep_goals:
                rep_content.append(f"Goals: {self.idp_data.rep_goals}")
            
            asp_content = []
            if self.idp_data.ultimate_goal:
                asp_content.append(f"Ultimate: {self.idp_data.ultimate_goal}")
            if self.idp_data.season_goal:
                asp_content.append(f"This Season: {self.idp_data.season_goal}")
            if self.idp_data.success_criteria:
                asp_content.append(f"Success: {self.idp_data.success_criteria}")
            
            table1.rows[1].cells[0].text = "\n".join(club_content) if club_content else "Not completed"
            table1.rows[1].cells[1].text = "\n".join(rep_content) if rep_content else "Not completed"
            table1.rows[1].cells[2].text = "\n".join(asp_content) if asp_content else "Not completed"
            
            doc.add_paragraph()
            
            # Table 2: The How
            table2 = doc.add_table(rows=2, cols=3)
            table2.style = 'Table Grid'
            
            headers2 = ["Fitness", "Law", "Mental"]
            for i, header in enumerate(headers2):
                cell = table2.rows[0].cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Build comprehensive fitness content
            fitness_content = []
            if self.idp_data.fitness_level:
                fitness_content.append(f"Current: {self.idp_data.fitness_level}")
            if self.idp_data.bronco_done:
                fitness_content.append(f"Bronco: {self.idp_data.bronco_done}")
            if self.idp_data.bronco_time:
                fitness_content.append(f"Current time: {self.idp_data.bronco_time}")
            if self.idp_data.bronco_target:
                fitness_content.append(f"Target: {self.idp_data.bronco_target}")
            if self.idp_data.training_frequency:
                fitness_content.append(f"Training: {self.idp_data.training_frequency}")
            if self.idp_data.training_types:
                fitness_content.append(f"Types: {self.idp_data.training_types}")
            if self.idp_data.fitness_target:
                fitness_content.append(f"Goal: {self.idp_data.fitness_target}")
            if self.idp_data.fitness_obstacles:
                fitness_content.append(f"Obstacles: {self.idp_data.fitness_obstacles}")
            
            # Build comprehensive law content
            law_content = []
            if self.idp_data.law_confidence:
                law_content.append(f"Confidence: {self.idp_data.law_confidence}/5")
            if self.idp_data.law_weak_areas:
                law_content.append(f"Weak areas: {self.idp_data.law_weak_areas}")
            if self.idp_data.law_consistency:
                law_content.append(f"Consistency: {self.idp_data.law_consistency}")
            if self.idp_data.law_plan:
                law_content.append(f"Plan: {self.idp_data.law_plan}")
            
            # Build comprehensive mental content
            mental_content = []
            if self.idp_data.mental_strengths:
                mental_content.append(f"Strengths: {self.idp_data.mental_strengths}")
            if self.idp_data.mental_challenges:
                mental_content.append(f"Challenges: {self.idp_data.mental_challenges}")
            if self.idp_data.mental_plan:
                mental_content.append(f"Plan: {self.idp_data.mental_plan}")
            
            table2.rows[1].cells[0].text = "\n".join(fitness_content) if fitness_content else "Not completed"
            table2.rows[1].cells[1].text = "\n".join(law_content) if law_content else "Not completed"
            table2.rows[1].cells[2].text = "\n".join(mental_content) if mental_content else "Not completed"
            
            doc.add_paragraph()
            
            # Table 3: Key Focus Areas
            table3 = doc.add_table(rows=2, cols=3)
            table3.style = 'Table Grid'
            
            headers3 = ["Focus Area 1", "Focus Area 2", "Focus Area 3"]
            for i, header in enumerate(headers3):
                cell = table3.rows[0].cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Build focus area 1 content
            focus1_content = []
            if self.idp_data.focus1_category:
                focus1_content.append(f"Category: {self.idp_data.focus1_category}")
            if self.idp_data.focus1_area:
                focus1_content.append(f"Area: {self.idp_data.focus1_area}")
            if self.idp_data.focus1_goal:
                focus1_content.append(f"Goal: {self.idp_data.focus1_goal}")
            if self.idp_data.focus1_how:
                focus1_content.append(f"How: {self.idp_data.focus1_how}")
            if self.idp_data.focus1_track:
                focus1_content.append(f"Track: {self.idp_data.focus1_track}")
            
            # Build focus area 2 content
            focus2_content = []
            if self.idp_data.focus2_category:
                focus2_content.append(f"Category: {self.idp_data.focus2_category}")
            if self.idp_data.focus2_area:
                focus2_content.append(f"Area: {self.idp_data.focus2_area}")
            if self.idp_data.focus2_goal:
                focus2_content.append(f"Goal: {self.idp_data.focus2_goal}")
            if self.idp_data.focus2_how:
                focus2_content.append(f"How: {self.idp_data.focus2_how}")
            if self.idp_data.focus2_track:
                focus2_content.append(f"Track: {self.idp_data.focus2_track}")
            
            # Build focus area 3 content
            focus3_content = []
            if self.idp_data.focus3_category:
                focus3_content.append(f"Category: {self.idp_data.focus3_category}")
            if self.idp_data.focus3_area:
                focus3_content.append(f"Area: {self.idp_data.focus3_area}")
            if self.idp_data.focus3_goal:
                focus3_content.append(f"Goal: {self.idp_data.focus3_goal}")
            if self.idp_data.focus3_how:
                focus3_content.append(f"How: {self.idp_data.focus3_how}")
            if self.idp_data.focus3_track:
                focus3_content.append(f"Track: {self.idp_data.focus3_track}")
            
            table3.rows[1].cells[0].text = "\n".join(focus1_content) if focus1_content else "Not completed"
            table3.rows[1].cells[1].text = "\n".join(focus2_content) if focus2_content else "Not completed"
            table3.rows[1].cells[2].text = "\n".join(focus3_content) if focus3_content else "Not completed"
            
            doc.save(filename)
            messagebox.showinfo("Success", f"IDP exported to:\n{filename}")
            
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export:\n{str(e)}")


def show_splash_screen(parent):
    """Display splash screen while app loads"""
    splash = tk.Toplevel(parent)
    splash.title("")
    splash.overrideredirect(True)  # Remove window decorations
    
    width = 500
    height = 300
    x = (splash.winfo_screenwidth() // 2) - (width // 2)
    y = (splash.winfo_screenheight() // 2) - (height // 2)
    splash.geometry(f'{width}x{height}+{x}+{y}')
    
    # Main frame
    frame = tk.Frame(splash, bg="#1976D2", relief=tk.RAISED, borderwidth=2)
    frame.pack(fill=tk.BOTH, expand=True)
    
    # App title
    tk.Label(frame, text="Rugby Referee Review System", 
            font=("Segoe UI", 20, "bold"), bg="#1976D2", fg="white").pack(pady=(40, 10))
    
    # Alpha Release badge
    alpha_frame = tk.Frame(frame, bg="#FF9800", relief=tk.RAISED, bd=1)
    alpha_frame.pack(pady=5)
    tk.Label(alpha_frame, text="ALPHA RELEASE 2", 
            font=("Segoe UI", 9, "bold"), bg="#FF9800", fg="white",
            padx=15, pady=3).pack()
    
    # Version
    tk.Label(frame, text=f"Version {__version__}", 
            font=("Segoe UI", 11), bg="#1976D2", fg="#E3F2FD").pack(pady=5)
    
    # Copyright
    tk.Label(frame, text="", bg="#1976D2").pack(pady=10)  # Spacer
    tk.Label(frame, text="Copyright © 2025 Andrew Clarkson", 
            font=("Segoe UI", 10), bg="#1976D2", fg="white").pack(pady=5)
    
    # CRRDF
    tk.Label(frame, text="Based on CRRDF Framework", 
            font=("Segoe UI", 9, "italic"), bg="#1976D2", fg="#E3F2FD").pack(pady=2)
    
    # Loading message
    tk.Label(frame, text="", bg="#1976D2").pack(pady=15)  # Spacer
    loading_label = tk.Label(frame, text="Loading...", 
            font=("Segoe UI", 10), bg="#1976D2", fg="#E3F2FD")
    loading_label.pack(pady=10)
    
    # Progress bar
    progress = ttk.Progressbar(frame, mode='indeterminate', length=300)
    progress.pack(pady=10)
    progress.start(10)
    
    splash.update()
    return splash


def main():
    """Main entry point"""
    # Create main window first (hidden)
    if THEME_AVAILABLE:
        root = ttk.Window(themename="cosmo")
    else:
        root = tk.Tk()
    
    root.withdraw()  # Hide main window initially
    
    # Show splash screen
    splash = show_splash_screen(root)
    
    # Simulate loading (in real app, this is where you'd load resources)
    import time
    for i in range(20):  # 2 seconds total
        time.sleep(0.1)
        splash.update()
    
    # Close splash and show main app
    splash.destroy()
    root.deiconify()  # Show main window
    
    app = CRRDFReviewApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()